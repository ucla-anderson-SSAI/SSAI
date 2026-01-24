"""
Generate Light Cases Word Documents for MGMT298D
UCLA Anderson - Science and Strategy of AI

Creates one Word document per week containing 3 light cases each.
Light cases are brief (5-minute) discussion prompts for mid-lecture breaks.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# UCLA Colors (RGB)
UCLA_BLUE = RGBColor(0x27, 0x74, 0xAE)
UCLA_DARK_BLUE = RGBColor(0x00, 0x55, 0x87)
UCLA_GOLD = RGBColor(0xFF, 0xD1, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)


def add_horizontal_line(paragraph, color_hex='FFD100'):
    """Add a horizontal line after a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_light_cases_document(filename, week_num, week_topic, cases):
    """Create a Word document containing all light cases for a week"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("MGMT298D: Science and Strategy of AI | UCLA Anderson")
    run.font.size = Pt(10)
    run.font.color.rgb = UCLA_BLUE
    run.font.name = 'Arial'

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Week {week_num} Light Cases")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(week_topic)
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Arial'

    # Instruction
    instruction = doc.add_paragraph()
    instruction.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = instruction.add_run("Brief discussion prompts for mid-lecture breaks (5 minutes each)")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = LIGHT_GRAY
    run.font.name = 'Arial'
    instruction.space_after = Pt(6)

    add_horizontal_line(instruction, 'FFD100')

    # Add each light case
    for i, case in enumerate(cases, 1):
        # Case header
        case_header = doc.add_paragraph()
        run = case_header.add_run(f"Light Case {i}: {case['title']}")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = UCLA_BLUE
        run.font.name = 'Arial'
        case_header.space_before = Pt(14)
        case_header.space_after = Pt(6)

        # Setup section
        setup_label = doc.add_paragraph()
        run = setup_label.add_run("Setup (1 minute)")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_DARK_BLUE
        run.font.name = 'Arial'
        setup_label.space_after = Pt(3)

        setup_text = doc.add_paragraph()
        run = setup_text.add_run(case['setup'])
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Arial'
        setup_text.paragraph_format.left_indent = Inches(0.15)
        setup_text.space_after = Pt(8)

        # Discussion prompt
        prompt_label = doc.add_paragraph()
        run = prompt_label.add_run("Discussion Prompt (3 minutes)")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_DARK_BLUE
        run.font.name = 'Arial'
        prompt_label.space_after = Pt(3)

        prompt_text = doc.add_paragraph()
        run = prompt_text.add_run(f'"{case["prompt"]}"')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Arial'
        prompt_text.paragraph_format.left_indent = Inches(0.15)
        prompt_text.space_after = Pt(8)

        # Teaching notes
        notes_label = doc.add_paragraph()
        run = notes_label.add_run("Key Points to Draw Out")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_DARK_BLUE
        run.font.name = 'Arial'
        notes_label.space_after = Pt(3)

        for note in case['notes']:
            note_para = doc.add_paragraph()
            run = note_para.add_run(f"  {note}")
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
            note_para.paragraph_format.left_indent = Inches(0.15)
            note_para.space_after = Pt(2)

        # Connection to content
        connection_label = doc.add_paragraph()
        run = connection_label.add_run("Connection to Content (1 minute)")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_DARK_BLUE
        run.font.name = 'Arial'
        connection_label.space_before = Pt(6)
        connection_label.space_after = Pt(3)

        connection_text = doc.add_paragraph()
        run = connection_text.add_run(case['connection'])
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Arial'
        connection_text.paragraph_format.left_indent = Inches(0.15)
        connection_text.space_after = Pt(4)

        # Source
        if case.get('source'):
            source_para = doc.add_paragraph()
            run = source_para.add_run(f"Source: {case['source']}")
            run.font.size = Pt(8)
            run.font.color.rgb = LIGHT_GRAY
            run.font.name = 'Arial'
            source_para.paragraph_format.left_indent = Inches(0.15)
            source_para.space_after = Pt(4)

        # Divider between cases (except last)
        if i < len(cases):
            divider = doc.add_paragraph()
            divider.space_before = Pt(4)
            add_horizontal_line(divider, 'CCCCCC')

    doc.save(filename)
    print(f"Created: {filename}")


# ============================================
# LIGHT CASES DATA BY WEEK
# ============================================

all_light_cases = {
    1: {
        "topic": "Linear Regression with Feature Engineering",
        "cases": [
            {
                "title": "Walmart Demand Forecasting",
                "setup": "Walmart manages millions of SKUs across thousands of stores. They use regression-based models to predict demand at the store-SKU-day level. The results: 16% reduction in stockouts and 10% improvement in inventory turnover.",
                "prompt": "Why might even a 1% improvement in forecast accuracy translate to hundreds of millions in savings for Walmart?",
                "notes": [
                    "Scale amplifies small improvements (millions of SKUs x thousands of stores x 365 days)",
                    "Stockouts = lost sales (customer goes to competitor)",
                    "Overstock = markdowns, spoilage, storage costs",
                    "Inventory ties up capital that could be invested elsewhere"
                ],
                "connection": "This is why the evaluation metrics we'll learn today-MAE and MSE-matter so much. A model that's 'pretty good' might leave millions on the table.",
                "source": "https://www.supplychaindive.com/news/4-walmart-supply-chain-ai-uses/760891/"
            },
            {
                "title": "Progressive Snapshot Insurance",
                "setup": "Progressive's Snapshot program puts a device in your car that tracks driving behavior-hard braking, time of day, miles driven. Their ML models use this data to price risk 9% more accurately than traditional factors. Safe drivers get average discounts of $130 after six months.",
                "prompt": "What are the privacy trade-offs of telematics-based insurance? Who might be disadvantaged by this approach?",
                "notes": [
                    "Trade-off: Privacy for savings (is $130 worth it?)",
                    "Winners: Safe drivers, people with predictable schedules",
                    "Losers: People who drive late at night (shift workers), long commuters",
                    "Equity concern: Low-income workers more likely to have variable schedules"
                ],
                "connection": "Progressive is using regression to predict claims from driving features. The coefficients they learn tell them: 'Hard braking at night is worth X more risk than daytime highway driving.' That's the interpretability we'll explore today.",
                "source": "https://www.progressive.com/auto/discounts/snapshot/"
            },
            {
                "title": "Uber Surge Pricing",
                "setup": "Uber uses ML to predict surge pricing in real-time, analyzing weather, events, historical patterns, traffic, and supply/demand balance across millions of rides. The model updates continuously as conditions change.",
                "prompt": "Uber's surge pricing is essentially a prediction problem-forecasting demand and supply imbalance. Why do consumers find it frustrating even if it's economically efficient?",
                "notes": [
                    "Fairness perception: 'Price gouging' during emergencies",
                    "Unpredictability: Hard to budget when prices vary 3x",
                    "Asymmetric information: Uber knows more than riders",
                    "Loss aversion: The 'normal' price feels like the fair price"
                ],
                "connection": "This shows that technical accuracy isn't everything. Zillow's model might have been accurate on average but wrong in exactly the situations where it mattered most. We'll talk about how to evaluate models not just on average error, but on the distribution of errors.",
                "source": "https://www.uber.com/blog/reinforcement-learning-for-modeling-marketplace-balance/"
            }
        ]
    },
    2: {
        "topic": "Tree-Based Predictions",
        "cases": [
            {
                "title": "Airbnb Search Ranking",
                "setup": "Airbnb's search ranking system determines which listings appear first-a decision worth millions in bookings. Their original system used gradient boosted decision trees (GBDT) trained on booking conversion. Only after years of refinement did they move to neural networks. Their VP of Engineering said: 'We started with GBDT because it was interpretable, fast to train, and good enough. Neural nets came later when we needed to incorporate images and text.'",
                "prompt": "Why might a company start with simpler tree-based models before investing in deep learning? What does 'good enough' mean when you're already making millions?",
                "notes": [
                    "Interpretability: Trees let you debug why bad listings rank high",
                    "Iteration speed: GBDTs train in minutes; neural nets take hours/days",
                    "Diminishing returns: Going from 70% to 90% accuracy is more valuable than 90% to 92%",
                    "Organizational readiness: Do you have the infrastructure for neural nets?",
                    "Technical debt: Simpler models are easier to maintain"
                ],
                "connection": "This shows a key principle: start simple, add complexity only when needed. The gradient boosting models we'll learn today are often the 'final answer' for structured data-not a stepping stone to deep learning.",
                "source": "https://medium.com/airbnb-engineering/machine-learning-powered-search-ranking-of-airbnb-experiences-110b4b1a0789"
            },
            {
                "title": "IBM Employee Attrition",
                "setup": "IBM developed a tree-based model to predict which employees were likely to resign, claiming 95% accuracy in identifying 'flight risks.' Managers received alerts about at-risk team members, allowing proactive retention efforts. IBM reported saving $300 million in retention costs. But the system raised immediate concerns: Should employees be scored on their likelihood to quit? What happens when managers treat predicted 'flight risks' differently?",
                "prompt": "If your employer could predict with 95% accuracy that you were about to quit, how would you want them to use that information? How might prediction become self-fulfilling prophecy?",
                "notes": [
                    "Positive uses: Better managers, proactive career conversations, fixing systemic issues",
                    "Negative uses: Withholding opportunities from 'flight risks,' creating hostile environment",
                    "Self-fulfilling prophecy: Treating someone as a flight risk may make them leave",
                    "Privacy concerns: What data goes into the model? Performance reviews? Emails?",
                    "Power asymmetry: Employer knows your intentions; you don't know you've been flagged"
                ],
                "connection": "This case shows that feature importance in trees reveals what the model 'thinks' matters. If the attrition model weights 'didn't get promoted' heavily, that's feedback about your organization-not just about individual employees.",
                "source": "https://www.cnbc.com/2019/04/03/ibm-ai-can-predict-with-95-percent-accuracy-which-employees-will-quit.html"
            },
            {
                "title": "First Hawaiian Bank",
                "setup": "First Hawaiian Bank, the oldest and largest bank in Hawaii, implemented tree-based ML models for loan decisioning in 2019. The results: a 13x increase in automatically approved loans and 50% reduction in decision time. Loan officers now focus on complex cases while routine applications are decided instantly. Customer satisfaction scores increased as approval times dropped from days to seconds.",
                "prompt": "A 13x increase in automated decisions sounds impressive, but what safeguards would you want before letting a model approve loans without human review?",
                "notes": [
                    "Threshold selection: What confidence level triggers auto-approval vs. human review?",
                    "Monitoring: Track approval rates, default rates, customer complaints by segment",
                    "Override capability: Can loan officers reverse model decisions?",
                    "Audit trail: If a loan defaults, can you explain why it was approved?",
                    "Fairness testing: Are approval rates equal across demographics?"
                ],
                "connection": "First Hawaiian shows the operational benefits of interpretable models. When regulators ask 'Why did you approve this loan?', a decision tree can provide the exact path. That explainability enables automation in regulated industries.",
                "source": "https://www.zest.ai/learn/success_stories/first-hawaiian-bank/"
            }
        ]
    },
    3: {
        "topic": "Clustering & Collaborative Filtering",
        "cases": [
            {
                "title": "Stitch Fix Style Profiles",
                "setup": "Stitch Fix sends personalized clothing boxes to customers based on 90+ data points collected through an onboarding quiz and ongoing feedback. CEO Katrina Lake said: 'Data science isn't a department at Stitch Fix-it's the foundation of the company.' Their algorithms cluster customers by style preferences, body characteristics, and budget, then match them with inventory. Human stylists review algorithmic recommendations before shipping, but the algorithm does the heavy lifting of narrowing millions of items to a shortlist.",
                "prompt": "Stitch Fix collects 90+ data points but starts with explicit customer preferences, not just behavioral data. When is asking customers what they want more valuable than inferring from their behavior?",
                "notes": [
                    "Behavioral data shows what they did; surveys show what they'd like",
                    "Cold start: New customers have no behavioral history",
                    "Preference vs. availability: Customers may want things they've never found",
                    "Trust building: Asking preferences makes customers feel heard",
                    "Combination is best: Initialize with stated preferences, refine with behavior"
                ],
                "connection": "Stitch Fix shows how clustering and collaborative filtering work together. Customers are clustered by style; inventory is matched via collaborative filtering ('customers like you loved this item'). The algorithm narrows the search; humans add judgment.",
                "source": "https://algorithms-tour.stitchfix.com/"
            },
            {
                "title": "Spotify Discover Weekly",
                "setup": "Discover Weekly is a personalized playlist Spotify generates for each user every Monday-40 million users receive a unique 30-song playlist tailored to their taste. It was born from a Hack Week project in 2015 and became one of Spotify's most beloved features. The algorithm combines collaborative filtering ('users with similar tastes liked these songs') with content-based filtering ('songs with similar audio features'). Users have discovered artists they'd never have found otherwise-and those artists have gained millions of new listeners.",
                "prompt": "Discover Weekly combines 'what similar users like' with 'what sounds similar.' Why might combining these approaches work better than either alone?",
                "notes": [
                    "Collaborative filtering: Finds non-obvious connections (jazz lovers who also like electronic)",
                    "Content-based: Maintains musical coherence (similar tempo, key, energy)",
                    "Combined: Discoverable but relevant-not too similar, not too random",
                    "Serendipity vs. relevance: Pure similarity is boring; pure novelty is random",
                    "The magic: Feels like a friend who knows you made it"
                ],
                "connection": "Discover Weekly is a masterclass in collaborative filtering at scale. 40 million personalized playlists, refreshed weekly, all powered by the matrix factorization techniques we'll learn today.",
                "source": "https://www.music-tomorrow.com/blog/how-spotify-recommendation-system-works-complete-guide"
            },
            {
                "title": "JPMorgan Customer Segmentation",
                "setup": "JPMorgan's wealth management division used clustering to identify distinct client segments beyond traditional 'high net worth' classifications. By analyzing investment behavior, risk tolerance, and life stage patterns, they identified segments like 'conservative accumulators,' 'aggressive traders,' and 'passive inheritors.' Tailoring advice and product recommendations to each segment increased wealth management sales by 20%. The insight: A 45-year-old entrepreneur and a 45-year-old surgeon with identical assets need very different financial advice.",
                "prompt": "Two clients with the same net worth might belong to different segments based on behavior. What behaviors would you cluster on for wealth management, and why?",
                "notes": [
                    "Risk behavior: Trading frequency, asset allocation, reaction to market drops",
                    "Life stage: Accumulating vs. preserving vs. distributing wealth",
                    "Goals: Retirement, legacy, lifestyle maintenance",
                    "Engagement: Self-directed vs. advice-seeking",
                    "Source of wealth: Earned vs. inherited vs. entrepreneurial exit"
                ],
                "connection": "This shows how clustering drives personalization in high-stakes advice. The segment someone belongs to determines which products are offered, which advisors are matched, and how communications are framed. Same clustering algorithm, very different application from Spotify.",
                "source": "https://digitaldefynd.com/IQ/jp-morgan-using-ai-case-study/"
            }
        ]
    },
    4: {
        "topic": "Reinforcement Learning",
        "cases": [
            {
                "title": "Netflix Thumbnail Personalization",
                "setup": "Netflix has multiple thumbnail images for each show. The question: Which thumbnail should they show to which user? This is a multi-armed bandit problem-each thumbnail is an 'arm,' and clicks are the reward. Using Thompson Sampling, Netflix learns personalized thumbnail preferences in real-time. A user who loves action movies might see an explosion scene; someone who loves romance might see a tender moment. The same show, different visual marketing. Netflix estimated that personalized thumbnails increased engagement by 20-30%.",
                "prompt": "Thumbnail personalization is a bandit problem: each thumbnail is an arm, each click is a reward. Why is this simpler than full RL, and where would you draw the line between bandit problems and RL problems?",
                "notes": [
                    "Bandits: Actions don't affect future state (showing thumbnail A doesn't change tomorrow's options)",
                    "RL: Actions affect future state (matching a driver now affects future availability)",
                    "Bandits are 'stateless'-each decision is independent",
                    "RL is 'stateful'-current decisions affect future options",
                    "Netflix thumbnails: Your choice today doesn't change your choices tomorrow (bandit)"
                ],
                "connection": "Multi-armed bandits are RL's simpler cousin. They teach us the exploration-exploitation trade-off without the complexity of state. Thompson Sampling, which Netflix uses, is one of the most elegant solutions-we'll see how it works today.",
                "source": "https://medium.com/netflix-techblog/artwork-personalization-c589f074ad76"
            },
            {
                "title": "DoorDash Thompson Sampling",
                "setup": "DoorDash faces a matching problem: which Dasher (delivery driver) should accept which order? They use Thompson Sampling to learn which Dashers are likely to accept which types of orders, optimizing for acceptance rate and delivery speed. The system maintains a probability distribution for each Dasher-order type combination, updating with each acceptance or rejection. When an order comes in, it 'samples' from these distributions to estimate which Dasher is most likely to accept AND deliver quickly. This reduced rejected orders by 8% and improved delivery times.",
                "prompt": "DoorDash uses Thompson Sampling to match Dashers to orders. What's the exploration-exploitation trade-off here? When should the system try a new Dasher vs. go with a known reliable one?",
                "notes": [
                    "Exploitation: Send to the Dasher with highest historical acceptance rate",
                    "Exploration: Try new Dashers to learn their preferences (they might be better)",
                    "Thompson Sampling: Automatically balances by sampling from uncertainty",
                    "High uncertainty = more exploration; low uncertainty = more exploitation",
                    "Real-world stakes: An 'exploration' attempt that fails means a delayed order"
                ],
                "connection": "DoorDash shows Thompson Sampling in a real marketplace with real consequences. Unlike A/B testing, which requires equal allocation, Thompson Sampling learns to allocate more traffic to better options while still exploring. It's 'adaptive' experimentation.",
                "source": "https://careersatdoordash.com/blog/using-a-multi-armed-bandit-with-thompson-sampling-to-identify-responsive-dashers/"
            },
            {
                "title": "Meta Ad Bidding with AdLlama",
                "setup": "Meta runs billions of ad auctions daily, deciding which ads to show to which users. Their 'AdLlama' system uses reinforcement learning to optimize bidding strategies across campaigns. Rather than treating each auction independently, AdLlama considers long-term value: a user who clicks today might convert tomorrow; a user who sees an ad today might buy next week. In tests, AdLlama achieved 6.7% improvement in click-through rate compared to previous systems-worth billions in advertiser value at Meta's scale.",
                "prompt": "Meta's ad system optimizes for long-term value, not just immediate clicks. A click today might lead to a purchase next week. How does this make the RL problem harder than optimizing for immediate clicks?",
                "notes": [
                    "Immediate clicks: Clear, fast feedback-easy to learn",
                    "Long-term value: Conversion might happen days later, on a different device",
                    "Attribution problem: Did this ad cause the purchase, or would they have bought anyway?",
                    "Delayed rewards: Must connect today's action to next week's outcome",
                    "Exploration is expensive: Every experimental ad shown is revenue at risk"
                ],
                "connection": "Ad optimization is RL at internet scale: billions of decisions, delayed rewards, and massive financial stakes. The techniques Meta uses-policy gradients, value estimation, credit assignment-are the same ones we'll learn today, just at a scale few can match.",
                "source": "https://ai.meta.com/research/publications/learning-to-bid-and-rank-together-in-recommendation-systems/"
            }
        ]
    },
    5: {
        "topic": "Neural Networks",
        "cases": [
            {
                "title": "DeepMind at National Grid",
                "setup": "Following success at Google data centers (40% cooling energy reduction), DeepMind partnered with the UK National Grid in 2022. The goal: use neural networks to predict electricity demand and optimize the entire British power grid. If successful, the system could reduce national energy consumption by 10%-equivalent to removing millions of cars from the road.",
                "prompt": "What's different about optimizing a national power grid versus a data center? What new challenges arise when AI systems make recommendations that affect millions of people?",
                "notes": [
                    "Scale: Millions of consumers with diverse needs vs. homogeneous servers",
                    "Stakes: Grid failure affects hospitals, emergency services, vulnerable populations",
                    "Accountability: Who's responsible when the AI is wrong?",
                    "Equity: Optimization might favor efficient areas over underserved ones",
                    "Transparency: Citizens may demand to know how AI affects their power"
                ],
                "connection": "This case shows neural networks scaling from controlled environments to messy real-world systems. The architecture stays similar-the challenges are data quality, uncertainty, and governance. We'll see how neural networks handle uncertainty when we discuss output layers and confidence calibration.",
                "source": "https://deepmind.google/blog/strengthening-our-partnership-with-the-uk-government-to-support-prosperity-and-security-in-the-ai-era/"
            },
            {
                "title": "UPS ORION Route Optimization",
                "setup": "UPS delivers 20 million packages daily. Their ORION system (On-Road Integrated Optimization and Navigation) uses neural networks to optimize delivery routes, considering traffic, weather, package priority, and driver constraints. The system evaluates 200,000 routing options per driver per day and has saved UPS 100 million miles annually-about $400 million in fuel costs.",
                "prompt": "ORION sometimes recommends routes that seem illogical to experienced drivers-like driving past a delivery location to return later. How do you balance algorithmic optimization against human intuition and experience?",
                "notes": [
                    "Global vs. local optimization: What seems inefficient locally might be optimal globally",
                    "Driver buy-in: Experienced drivers initially resisted 'dumb' AI suggestions",
                    "Edge cases: AI doesn't know about construction, dog attacks, customer preferences",
                    "Hybrid approach: UPS lets drivers override within limits, learns from overrides",
                    "Trust calibration: Drivers learned when to trust and when to override"
                ],
                "connection": "ORION shows neural networks tackling combinatorial optimization-finding the best path among trillions of possibilities. We'll see today how hidden layers allow networks to learn complex, non-linear relationships that traditional optimization struggles with.",
                "source": "https://www.informs.org/Impact/O.R.-Analytics-Success-Stories/Optimizing-Delivery-Routes"
            },
            {
                "title": "Walmart Inventory Neural Networks",
                "setup": "Walmart uses neural network-based demand forecasting across 4,700 US stores and millions of SKUs. The system predicts demand at the store-item-day level, incorporating weather, events, holidays, and local factors. Results: 16% reduction in stockouts and significant decrease in overstock markdowns.",
                "prompt": "Walmart's neural network must predict demand for millions of unique store-product-day combinations. Why might a neural network outperform traditional statistical methods at this scale?",
                "notes": [
                    "Shared learning: Neural network learns patterns that transfer across similar products",
                    "Feature learning: Automatically discovers relevant features from raw data",
                    "Non-linear relationships: Weather affects ice cream differently than soup",
                    "Scale efficiency: One model for millions of predictions vs. millions of individual models",
                    "But: Interpretability loss-hard to explain why prediction is what it is"
                ],
                "connection": "Walmart's system shows the power of neural networks for learning shared representations. The same network that predicts ice cream demand in Arizona learns patterns that help predict soda demand in Minnesota. This 'transfer' happens automatically through learned features-something we'll explore when we discuss hidden layer representations.",
                "source": "https://www.supplychaindive.com/news/walmart-grocery-AI-demand-operations/585424/"
            }
        ]
    },
    6: {
        "topic": "Convolutional Neural Networks",
        "cases": [
            {
                "title": "Second Spectrum NBA Tracking",
                "setup": "Second Spectrum (acquired by Genius Sports for $200 million) uses CNN-based computer vision to track every player and the ball in NBA games. The system powers real-time graphics you see on broadcasts-augmented reality overlays showing shooting probability, defensive coverage, and play diagrams. What once required manual annotation now happens in real-time, creating entirely new ways to visualize sports.",
                "prompt": "Second Spectrum's tracking enables 'shot probability' graphics that update in real-time as plays develop. How might showing probabilities during live broadcasts change how fans watch games? Could it change how coaches and players make decisions?",
                "notes": [
                    "Fan engagement: Makes complex basketball more accessible (like FiveThirtyEight for elections)",
                    "Coaching: Real-time feedback on defensive breakdowns, optimal shot selection",
                    "Player behavior: Might players avoid 'low probability' shots even when strategically correct?",
                    "Betting implications: Real-time probability = real-time betting markets",
                    "ESPN vs. sports books: Same data, very different uses"
                ],
                "connection": "This is CNNs operating in real-time on video-tracking tiny objects (basketball) and fast-moving people across different camera angles. We'll see how convolutional architectures can process temporal sequences when we discuss video understanding extensions.",
                "source": "https://pr.nba.com/nba-genius-sports-second-spectrum-expanded-partnership/"
            },
            {
                "title": "Nike Fit App",
                "setup": "Nike Fit uses smartphone cameras and CNNs to create 3D scans of customers' feet, recommending the right shoe size for each Nike model. The problem Nike is solving: shoe returns cost retailers $50 billion annually, with sizing being the top reason. Nike's CNN measures feet to within 2mm accuracy-better than most in-store measuring devices.",
                "prompt": "Nike Fit could reduce returns significantly, but it also gives Nike data about millions of feet. What other applications might Nike find for 3D foot scanning data? What competitive advantages does this create?",
                "notes": [
                    "Product design: Actual foot shape data could inform shoe design",
                    "Personalization: Custom-fit shoes at mass-market prices",
                    "Health insights: Changes in foot shape could indicate health issues",
                    "Competitive moat: Data asset that competitors can't easily replicate",
                    "Platform potential: Could Nike license the technology?"
                ],
                "connection": "Nike Fit is 3D reconstruction from 2D images-using CNNs to infer depth and shape from smartphone photos. We'll see how convolutional architectures can learn to estimate 3D structure from 2D input, similar to what Tesla's Autopilot does for driving.",
                "source": "https://www.cnbc.com/2019/05/08/nikes-app-is-getting-a-new-sizing-feature-heres-how-it-will-work.html"
            },
            {
                "title": "Medical Imaging AI",
                "setup": "75% of FDA-approved AI medical devices are in radiology-reading X-rays, CT scans, and MRIs. Companies like Aidoc, Viz.ai, and Paige use CNNs to detect tumors, strokes, and fractures, often faster and more consistently than human radiologists. Viz.ai's stroke detection system alerts specialists within 6 minutes of a brain scan-in a condition where 'time is brain' and every minute matters.",
                "prompt": "Medical imaging AI can match or exceed radiologist accuracy on many tasks. How should we think about AI diagnosis: as a replacement for doctors, as a tool for doctors, or something else?",
                "notes": [
                    "Augmentation vs. replacement: Most deployments augment (second opinion, triage)",
                    "Liability: Who's responsible when AI misses a tumor? The company? The doctor?",
                    "Workflow integration: Best AI is invisible-fits into existing processes",
                    "Access expansion: Could bring specialist-level diagnosis to underserved areas",
                    "Remaining gap: AI can find patterns but can't integrate with patient history, context"
                ],
                "connection": "Medical imaging is perhaps the highest-stakes CNN application: life-or-death decisions based on pattern recognition. The same convolutional features that detect cats' ears can detect tumor margins-it's all about learning discriminative visual patterns.",
                "source": "https://www.fda.gov/news-events/press-announcements/fda-permits-marketing-clinical-decision-support-software-alerting-providers-potential-stroke"
            }
        ]
    },
    7: {
        "topic": "Word Embeddings & Transformers",
        "cases": [
            {
                "title": "NYT Dynamic Paywall",
                "setup": "The New York Times uses causal machine learning to optimize its paywall. Rather than showing every reader the same 'Subscribe after 5 articles' message, they predict each reader's price sensitivity and subscription likelihood. High-value readers who are about to subscribe see fewer barriers; casual readers see more prompts. The result: significant lift in conversions without revealing the exact paywall rules.",
                "prompt": "If two readers click on the same article, one might hit a paywall while the other reads freely. Is this fair? What are the ethical considerations of personalized pricing and access?",
                "notes": [
                    "Price discrimination is legal and common (student discounts, coupons)",
                    "But algorithmic discrimination feels different-users don't know the rules",
                    "Could exacerbate information inequality: Wealthy readers pay less",
                    "Tension between business optimization and equitable access",
                    "The NYT is a public interest institution-does that change the calculation?"
                ],
                "connection": "This personalization requires understanding reader behavior and content-exactly what embeddings and transformers enable. The NYT's model predicts not just 'will they click?' but 'what content do they value?' That requires semantic understanding.",
                "source": "https://venturebeat.com/ai/how-machine-learning-helps-the-new-york-times-power-its-paywall"
            },
            {
                "title": "Google Translate Evolution",
                "setup": "In 2016, Google switched Google Translate from phrase-based statistical translation to neural machine translation. Overnight, error rates dropped 55-85% for many language pairs. Sentences that had been incomprehensible became fluent. The transformer architecture, introduced in 2017, made another leap possible-understanding context across entire paragraphs.",
                "prompt": "Google Translate now handles 100+ languages, including many with limited training data. How might translation quality vary across languages, and what are the implications for global communication?",
                "notes": [
                    "High-resource languages (English, Spanish, French) have excellent quality",
                    "Low-resource languages (many African, Indigenous languages) still struggle",
                    "This creates a digital divide: Some cultures well-represented, others not",
                    "Biases in training data propagate (gender defaults, cultural assumptions)",
                    "Professional translation still needed for high-stakes content (legal, medical)"
                ],
                "connection": "Neural machine translation was the first killer app for sequence-to-sequence attention models. The key insight: words have different meanings in different contexts. 'Bank' means something different in 'river bank' vs. 'bank account.' Attention mechanisms learn to disambiguate by looking at surrounding words.",
                "source": "https://research.google/blog/a-neural-network-for-machine-translation-at-production-scale/"
            },
            {
                "title": "Content Moderation at Scale",
                "setup": "Meta reports that 95% of hate speech removed from Facebook and Instagram is detected by AI before any human reports it. The platform removes millions of pieces of content daily across 50+ languages. But the system isn't perfect-appeals often succeed, and false positives affect legitimate speech.",
                "prompt": "AI content moderation must make split-second decisions about context-dependent content: satire, reclaimed language, news reporting on hate groups. How should platforms balance automation efficiency against the nuance these cases require?",
                "notes": [
                    "Scale makes human review impossible (billions of posts daily)",
                    "But nuance requires human judgment (satire, context, intent)",
                    "Errors have real consequences: Wrongly banned users, undetected harm",
                    "Cultural context varies: Acceptable in one country, offensive in another",
                    "Transparency: Users don't understand why content was removed"
                ],
                "connection": "This is transformer NLU at the hardest level: understanding not just what words mean, but the speaker's intent, cultural context, and downstream impact. The model must recognize 'I hate you' as different in a breakup song vs. a threat. These systems are remarkably good-and still far from sufficient.",
                "source": "https://transparency.meta.com/enforcement/"
            }
        ]
    },
    8: {
        "topic": "Generative AI APIs & AI Agents",
        "cases": [
            {
                "title": "Financial Times AI Strategy",
                "setup": "The Financial Times made an unusual choice in the AI wars: They licensed their content to OpenAI while simultaneously building 'Ask FT'-their own AI product that answers subscriber questions using FT journalism. Unlike newspapers suing AI companies, the FT chose to participate. CEO John Ridding explained: 'If AI is going to summarize news anyway, we'd rather shape how our journalism is used.'",
                "prompt": "The FT is both a supplier to AI companies (licensing content) and a competitor (building Ask FT). How should media companies navigate being simultaneously partner and rival to AI platforms?",
                "notes": [
                    "Licensing creates revenue but also trains competitors",
                    "Building own AI maintains customer relationship but requires investment",
                    "'Co-opetition': Common in tech (Apple uses Google search while competing)",
                    "The real question: Where is the value created and captured?",
                    "User trust: FT readers trust FT journalism, not generic AI"
                ],
                "connection": "The FT's strategy is a business response to the technology we're learning today: retrieval-augmented generation. Ask FT retrieves from FT articles, ensuring responses are grounded in quality journalism. The technical architecture enables the business model.",
                "source": "https://openai.com/index/content-partnership-with-financial-times/"
            },
            {
                "title": "Enterprise LLM Market Explosion",
                "setup": "The enterprise market for large language models is projected to grow from $6.7 billion in 2023 to $71.1 billion by 2030-a 41% compound annual growth rate. This makes it one of the fastest-growing enterprise software categories in history. Every major cloud provider now offers LLM APIs, and startups are raising billions to build on top of them.",
                "prompt": "If everyone has access to the same LLM APIs (OpenAI, Anthropic, Google), where does competitive advantage come from? What differentiates companies using the same underlying models?",
                "notes": [
                    "Proprietary data: Fine-tuning on your unique information",
                    "Integration: How well AI connects to existing workflows",
                    "Prompt engineering: Some companies get better results from same models",
                    "Speed of iteration: Rapid experimentation and deployment",
                    "Domain expertise: Understanding what questions to ask",
                    "The model is table stakes; the application is the differentiator"
                ],
                "connection": "This market growth is why we're learning LLM APIs today. These aren't research tools anymore-they're infrastructure. The question isn't whether your company will use LLMs, but how effectively. The skills you're learning will be essential across industries.",
                "source": "https://www.grandviewresearch.com/industry-analysis/enterprise-llm-market-report"
            },
            {
                "title": "ChatGPT Enterprise Deployment",
                "setup": "In 2023, OpenAI launched ChatGPT Enterprise with features demanded by corporate customers: no training on company data, SOC 2 compliance, admin controls, analytics, and unlimited high-speed access. Within months, companies like PwC, Klarna, and Canva deployed it to tens of thousands of employees. But deployment revealed new challenges: Shadow AI (employees using personal accounts), prompt injection attacks, and confidential data accidentally shared with AI.",
                "prompt": "Your company just deployed ChatGPT Enterprise to all employees. What governance policies and guardrails would you implement to maximize benefit while managing risk?",
                "notes": [
                    "Data classification: What information can and can't go into AI",
                    "Training: Employees need to know capabilities and limitations",
                    "Monitoring: Are people using it effectively? Sharing sensitive data?",
                    "Shadow AI: Block personal accounts or embrace and redirect?",
                    "Compliance: Industry-specific regulations (healthcare, finance)",
                    "Change management: Cultural resistance from some employees"
                ],
                "connection": "Enterprise deployment is where everything we've learned meets organizational reality. The technical challenge-building prompts and integrations-is only half the battle. The other half is governance, change management, and measuring impact. This is why AI literacy matters for managers, not just engineers.",
                "source": "https://openai.com/index/introducing-chatgpt-enterprise/"
            }
        ]
    }
}


# ============================================
# GENERATE LIGHT CASES DOCUMENTS
# ============================================

os.makedirs("/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/light-cases", exist_ok=True)

for week_num, week_data in all_light_cases.items():
    filename = f"/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/light-cases/Week{week_num}_Light_Cases.docx"
    create_light_cases_document(
        filename,
        week_num,
        week_data['topic'],
        week_data['cases']
    )

print("\nLight cases documents generated successfully!")
print("Created one document per week (8 total) in the light-cases folder.")
