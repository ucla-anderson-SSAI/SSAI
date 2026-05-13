"""
Generate Word Document Case Study Handouts for MGMT298D
UCLA Anderson - Science and Strategy of AI

Creates:
1. Student handouts (.docx) - editable Word documents
2. Instructor notes (.docx) - with teaching guidance and key insights

All facts are drawn from cited sources to minimize hallucination.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# UCLA Colors (RGB)
UCLA_BLUE = RGBColor(0x27, 0x74, 0xAE)
UCLA_DARK_BLUE = RGBColor(0x00, 0x55, 0x87)
UCLA_GOLD = RGBColor(0xFF, 0xD1, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)


def add_horizontal_line(paragraph, color_hex='FFD100'):
    """Add a horizontal line after a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # line thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_student_handout(filename, week_num, week_topic, case_title, content_paragraphs, key_facts, questions, primary_source):
    """Create student-facing handout as Word document"""
    doc = Document()

    # Set margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Course Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("MGMT298D: Science and Strategy of AI | UCLA Anderson")
    run.font.size = Pt(10)
    run.font.color.rgb = UCLA_BLUE
    run.font.name = 'Arial'

    # Case Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(case_title)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'

    # Week Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Week {week_num}: {week_topic}")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Arial'

    # Gold divider line
    add_horizontal_line(subtitle, 'FFD100')

    # Section Header: The Situation
    section_header = doc.add_paragraph()
    run = section_header.add_run("The Situation")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    section_header.space_before = Pt(6)
    section_header.space_after = Pt(3)

    # Content paragraphs
    for para in content_paragraphs:
        p = doc.add_paragraph()
        # Check if it's a quote (starts with quotation marks)
        if para.startswith('"') or para.startswith('"'):
            run = p.add_run(para)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
        else:
            run = p.add_run(para)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Section Header: Discussion Questions
    section_header = doc.add_paragraph()
    run = section_header.add_run("Discussion Questions")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    section_header.space_before = Pt(6)
    section_header.space_after = Pt(3)

    # Questions
    for i, q in enumerate(questions, 1):
        # Question number
        q_num = doc.add_paragraph()
        run = q_num.add_run(f"Question {i}")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_BLUE
        run.font.name = 'Arial'
        q_num.space_before = Pt(4)
        q_num.space_after = Pt(1)

        # Question text
        q_text = doc.add_paragraph()
        run = q_text.add_run(q['question'])
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Arial'
        q_text.paragraph_format.left_indent = Inches(0.25)
        q_text.space_after = Pt(3)

    # Footer with source
    footer_line = doc.add_paragraph()
    footer_line.space_before = Pt(6)
    add_horizontal_line(footer_line, 'CCCCCC')

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Source: {primary_source}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Arial'

    doc.save(filename)
    print(f"Created: {filename}")


def create_instructor_notes(filename, week_num, week_topic, case_title, questions, key_insight, teaching_tips):
    """Create instructor-only teaching notes as Word document"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Instructor Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("INSTRUCTOR NOTES - DO NOT DISTRIBUTE")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run.font.name = 'Arial'

    # Case Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(case_title)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'

    # Week Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Week {week_num}: {week_topic}")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Arial'

    # Red divider line
    add_horizontal_line(subtitle, 'C62828')

    # Key Insight Section
    insight_header = doc.add_paragraph()
    run = insight_header.add_run("Key Insight (Goal of Discussion)")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run.font.name = 'Arial'
    insight_header.space_before = Pt(12)
    insight_header.space_after = Pt(6)

    insight_box = doc.add_paragraph()
    insight_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = insight_box.add_run(key_insight)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    insight_box.space_after = Pt(10)

    # Discussion Guide Header
    guide_header = doc.add_paragraph()
    run = guide_header.add_run("Discussion Guide")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    guide_header.space_before = Pt(10)
    guide_header.space_after = Pt(6)

    # Questions with teaching notes
    for i, q in enumerate(questions, 1):
        # Question
        q_para = doc.add_paragraph()
        run = q_para.add_run(f"Question {i}: {q['question']}")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_BLUE
        run.font.name = 'Arial'
        q_para.space_before = Pt(8)
        q_para.space_after = Pt(4)

        # Teaching Notes label
        notes_label = doc.add_paragraph()
        run = notes_label.add_run("Teaching Notes:")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Arial'
        notes_label.paragraph_format.left_indent = Inches(0.25)
        notes_label.space_after = Pt(2)

        # Teaching notes bullets
        for note in q['notes']:
            note_para = doc.add_paragraph()
            run = note_para.add_run(f"  {note}")
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
            note_para.paragraph_format.left_indent = Inches(0.25)
            note_para.space_after = Pt(2)

    # Additional Teaching Tips
    if teaching_tips:
        tips_header = doc.add_paragraph()
        run = tips_header.add_run("Additional Teaching Tips")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        run.font.name = 'Arial'
        tips_header.space_before = Pt(12)
        tips_header.space_after = Pt(6)

        for tip in teaching_tips:
            tip_para = doc.add_paragraph()
            run = tip_para.add_run(f"  {tip}")
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
            tip_para.paragraph_format.left_indent = Inches(0.25)
            tip_para.space_after = Pt(4)

    doc.save(filename)
    print(f"Created: {filename}")


# ============================================
# CASE DATA - With Expanded Source-Based Content
# ============================================

cases = [
    # WEEK 1A: ZILLOW
    {
        "week": 1,
        "topic": "Linear Regression",
        "id": "Week1A_Zillow",
        "title": "Zillow's $500M iBuying Failure",
        "content_paragraphs": [
            "In 2018, Zillow launched Zillow Offers, an ambitious program to use its famous Zestimate algorithm to make instant cash offers on homes. The idea seemed brilliant: leverage 15 years of home valuation data to buy homes, make light renovations, and resell quickly for profit. By late 2021, the company was on track to purchase 5,000 homes per month.",

            "Then everything collapsed. In Q3 2021-the three months that broke the business-Zillow purchased more houses than in the previous 18 months combined. The company bought 9,680 homes in that quarter alone, up from 3,805 in Q2. But they only sold 3,032 homes, far below expectations. The average gross profit per home sold was a loss of $80,771.",

            '"We have determined the unpredictability in forecasting home prices far exceeds what we anticipated and continuing to scale Zillow Offers would result in too much earnings and balance-sheet volatility," CEO Rich Barton told investors. He admitted the company realized it could not predict where home prices would be in six months "within a narrow margin of error."',

            "The market had started to cool in summer 2021, but Zillow's algorithm did not see the brake lights. According to real estate analyst Mike DelPrete: 'What happened is they ended up buying a lot of houses and overpaying for a lot of those houses because they did not see the market start to cool.' Stanford professor Amit Seru noted that Zillow made a critical strategic error: 'They were late entrants into this market and decided, among other things, to go into non-cookie-cutter homes, hoping their algorithmic valuation model was accurate.'",

            "In November 2021, Zillow announced it would shut down Zillow Offers entirely, take a $569 million write-down (about $30,000 per home in inventory), and lay off 2,000 employees-25% of its workforce. The stock plummeted, wiping out approximately $7.8 billion in market value within days. Meanwhile, competitors Opendoor and Offerpad-using AI models that detected the cooling market-continued operating.",
        ],
        "key_facts": [
            ("Zestimate launched", "2006 (15 years before iBuying)"),
            ("iBuying launched", "2018 (Phoenix and Las Vegas)"),
            ("Q3 2021 homes purchased", "9,680 (vs. 3,805 in Q2)"),
            ("Q3 2021 loss per home", "$80,771 average"),
            ("Total write-down", "$569 million (~$30K per home)"),
            ("Employees laid off", "2,000 (25% of workforce)"),
            ("Market cap lost", "~$7.8 billion in days"),
        ],
        "questions": [
            {
                "question": "Zillow's Zestimate had been running for 15 years before iBuying. What fundamentally changed when they moved from providing estimates to acting on them?",
                "notes": [
                    "Providing information vs. making decisions have different error tolerances",
                    "When you are just informing, being 5% off is fine-users adjust",
                    "When you are betting $500K per house, 5% off is catastrophic",
                    "The model did not get worse; the stakes changed"
                ]
            },
            {
                "question": "Zillow lost $80,771 per home on average in Q3 2021. What features would you include in a home price prediction model, and which would be hardest to capture?",
                "notes": [
                    "Easy: Square footage, bedrooms, location, lot size, age",
                    "Medium: School quality, crime rates, commute times",
                    "Hard: Neighborhood 'feel,' future development plans, market momentum",
                    "Impossible: What a specific buyer will pay (idiosyncratic preferences)"
                ]
            },
            {
                "question": "Competitors Opendoor and Offerpad detected the cooling market and adjusted. Why might Zillow's algorithm have missed these signals?",
                "notes": [
                    "Training data was pre-pandemic 'normal' market",
                    "Models assume the future resembles the past (stationarity assumption)",
                    "Aggressive market share strategy may have overridden model warnings",
                    "Non-cookie-cutter homes harder to value accurately"
                ]
            },
            {
                "question": "If you were the CEO, what guardrails would you have put in place before betting $1 billion on algorithmic home purchases?",
                "notes": [
                    "Human review for high-value purchases",
                    "Maximum bid limits (cap the downside)",
                    "Geographic diversification requirements",
                    "Real-time model monitoring for drift",
                    "Stress testing against historical volatility"
                ]
            }
        ],
        "key_insight": "The gap between 'good at prediction' and 'good enough to bet the company' is enormous. Prediction accuracy is not the same as decision quality.",
        "primary_source": "stanford.edu/insights/flip-flop-why-zillows-algorithmic-home-buying-venture-imploded",
        "teaching_tips": [
            "Start by asking who has used Zestimate-most students will have",
            "Emphasize this was not a 'bad model'-it was a misapplication of a good model",
            "Connect to the lecture: What MAE would have been acceptable for iBuying?",
            "This case sets up the theme: prediction does not equal decision quality"
        ]
    },

    # WEEK 1B: MONEYBALL
    {
        "week": 1,
        "topic": "Linear Regression",
        "id": "Week1B_Moneyball",
        "title": "Moneyball: How the Oakland A's Beat the Yankees with Data",
        "content_paragraphs": [
            "In 2002, the Oakland Athletics faced an impossible challenge. After losing MVP Jason Giambi (7-year, $120M contract), Johnny Damon (4-year, $31M), and closer Jason Isringhausen (4-year, $27M)-totaling $31.6 million in annual value-the cash-strapped A's had a payroll of just $41 million, less than a third of the New York Yankees' $125 million.",

            "General Manager Billy Beane, who became GM in 1997, turned to sabermetrics-a term coined by statistician Bill James for the empirical analysis of baseball. With assistant GM Paul DePodesta (a Harvard economics graduate), Beane began looking at the game in a fundamentally different way than traditional scouts.",

            "The key insight came from regression analysis: on-base percentage (OBP) and slugging percentage (SLG) were far better predictors of runs scored than the traditional statistics scouts valued-batting average, stolen bases, and RBIs. Players who drew walks were seen as 'passive' by traditional scouts, but the data showed walks were just as valuable as hits for getting on base. Beane believed that 'power could be developed, but patience at the plate and the ability to get on base could not.'",

            "The A's draft strategy also diverged from convention. They focused on college players over high school prospects because college players had more statistical data available, allowing for better performance projections. In the 2002 draft, they selected unconventional prospects like Nick Swisher and Joe Blanton based on statistical analysis-both became successful major leaguers.",

            "The results were remarkable. The 2002 Oakland A's won 103 games and became the first team in over 100 years of American League baseball to win 20 consecutive games-a record that still stands. They made the playoffs with roughly one-third the Yankees' payroll.",

            "The A's advantage, however, proved temporary. In the wake of Moneyball, teams including the Boston Red Sox, New York Yankees, and Cleveland Guardians hired full-time sabermetric analysts. By 2016, all 30 NBA teams and 26 of 30 MLB teams had dedicated analytics departments. The market inefficiency that Beane exploited had been competed away.",
        ],
        "key_facts": [
            ("A's payroll (2002)", "$41M (3rd lowest in MLB)"),
            ("Yankees payroll (2002)", "$125M (highest)"),
            ("Star players lost", "$31.6M annual value (Giambi, Damon, Isringhausen)"),
            ("A's wins (2002)", "103 games, first-round playoff exit"),
            ("Winning streak", "20 games (still AL record)"),
            ("Key discovery", "OBP + SLG predict runs better than batting average"),
            ("By 2016", "26 of 30 MLB teams had analytics staff"),
        ],
        "questions": [
            {
                "question": "Beane discovered that on-base percentage was 'undervalued' compared to batting average. What made OBP undervalued, and why did scouts miss this for decades?",
                "notes": [
                    "Batting average was traditional, visible, discussed by media",
                    "OBP included walks, which seemed 'passive' to scouts",
                    "Regression revealed OBP correlated more strongly with runs scored",
                    "Market inefficiency: scouts used intuition, not data"
                ]
            },
            {
                "question": "The A's advantage eroded as other teams adopted analytics. How do you maintain competitive advantage when your methods can be copied?",
                "notes": [
                    "First-mover advantage is temporary",
                    "Must continuously innovate (new data sources, better models)",
                    "Today's edge: proprietary data, faster iteration, better execution",
                    "Similar pattern in business: early AI adopters lose edge as tools commoditize"
                ]
            },
            {
                "question": "Baseball is data-rich with clear outcomes. What business domains have similar characteristics? Which don't?",
                "notes": [
                    "Similar: Finance (stock prices), retail (sales), digital marketing (clicks)",
                    "Different: B2B sales (few observations), healthcare (long feedback loops), creativity (subjective outcomes)",
                    "Key requirements: abundant data, clear outcome variable, fast feedback"
                ]
            },
            {
                "question": "The A's won 103 games but lost in the first round of playoffs. Does this represent success or failure of the Moneyball approach?",
                "notes": [
                    "Regular season: Large sample, skill dominates",
                    "Playoffs: Small sample, luck matters more",
                    "Optimizing for playoffs vs. regular season may require different strategies",
                    "The model worked-the sample size was just too small"
                ]
            }
        ],
        "key_insight": "Regression reveals which variables actually matter-often not the ones intuition suggests. But analytical advantages erode as competitors adopt similar methods.",
        "primary_source": "en.wikipedia.org/wiki/Moneyball:_The_Art_of_Winning_an_Unfair_Game",
        "teaching_tips": [
            "Ask how many students have seen the movie-use it as a touchpoint",
            "Emphasize this is about regression coefficients revealing truth",
            "Connect to lecture: This is what beta coefficients tell us",
            "Good discussion on competitive advantage and data moats"
        ]
    },

    # WEEK 1C: AIRBNB
    {
        "week": 1,
        "topic": "Linear Regression",
        "id": "Week1C_Airbnb",
        "title": "Airbnb Smart Pricing: Optimization at Scale",
        "content_paragraphs": [
            "Airbnb hosts face a challenging decision every day: price too high and the listing sits empty; price too low and money is left on the table. To solve this, Airbnb developed Smart Pricing-a machine learning system that suggests optimal price adjustments based on demand signals, updated multiple times daily.",

            "The system analyzes more than 70 factors to generate price recommendations. These include historical booking data, current market conditions, seasonal trends, local events, competitor pricing, and even flight data to predict demand spikes. The model also considers listing-specific features like location, amenities, and quality metrics derived from user engagement.",

            "Technically, Airbnb uses Random Forest models rather than simple regression because pricing relationships are not linear. Each decision tree in the forest learns different patterns, and the ensemble captures complex interactions between variables. The company also employs reinforcement learning to improve pricing decisions based on outcomes, and elasticity models to gauge how sensitive demand is to price changes.",

            "But Smart Pricing has drawn criticism. Reports suggest the algorithm 'favors calendar fill over revenue maximization'-optimizing occupancy metrics that benefit the platform more than individual hosts. As one analysis noted: 'New Airbnb hosts accept these recommendations without knowing' that the system may be optimizing for Airbnb's interests rather than host revenue.",

            "The system also has notable limitations. Big cities with thousands of listings provide ample training data, but rural areas with few listings produce less reliable recommendations. Unique properties like treehouses challenge models trained mainly on standard homes. And during unprecedented situations-like the pandemic-historical data becomes irrelevant.",

            "Third-party alternatives like PriceLabs, Beyond Pricing, and Wheelhouse have emerged, marketing themselves as optimizing host profits rather than bookings. These tools analyze multiple platforms (Airbnb, Vrbo, Booking.com), while Smart Pricing only sees Airbnb demand. Hosts must choose between convenience and potentially conflicting incentives.",
        ],
        "key_facts": [
            ("Factors analyzed", "70+ variables"),
            ("Update frequency", "Multiple times daily"),
            ("Model type", "Random Forest (not linear regression)"),
            ("Data sources", "Historical bookings, events, weather, competitors, flights"),
            ("Criticism", "May optimize occupancy over host revenue"),
            ("Limitation", "Works best in data-rich urban markets"),
            ("Alternatives", "PriceLabs, Beyond Pricing, Wheelhouse"),
        ],
        "questions": [
            {
                "question": "Smart Pricing analyzes 70+ factors. Why might Airbnb use Random Forest instead of linear regression for this problem?",
                "notes": [
                    "Pricing relationships are not linear",
                    "Interactions between variables matter (event + weekend + summer)",
                    "Random Forest captures non-linear patterns",
                    "Ensemble methods reduce overfitting"
                ]
            },
            {
                "question": "Critics say Smart Pricing optimizes occupancy over revenue. What is the difference, and why might Airbnb prefer one?",
                "notes": [
                    "Airbnb earns fees on every booking (wants volume)",
                    "Hosts want revenue (price times occupancy)",
                    "Lower prices increase occupancy but may decrease total revenue",
                    "Platform vs. participant incentive misalignment"
                ]
            },
            {
                "question": "The system works well in cities but struggles in rural areas. Why does data density matter for ML pricing?",
                "notes": [
                    "More data equals better pattern recognition",
                    "Sparse data equals unreliable predictions",
                    "Unique properties have few comparables",
                    "Cold start problem for new markets"
                ]
            },
            {
                "question": "How should Smart Pricing handle unprecedented situations like a pandemic that eliminates travel demand?",
                "notes": [
                    "Historical data becomes irrelevant",
                    "Model will be wrong in unpredictable ways",
                    "Need: Human override capability, anomaly detection, rapid retraining",
                    "Similar challenge to Zillow: models assume stationarity"
                ]
            }
        ],
        "key_insight": "Pricing optimization is a natural application for ML-clear outcome, abundant data, actionable predictions. But platform incentives may not align with user incentives.",
        "primary_source": "medium.com/airbnb-engineering/learning-market-dynamics-for-optimal-pricing-97cffbcc53e3",
        "teaching_tips": [
            "Good for discussing platform economics and incentive alignment",
            "Ask students: Would you use Smart Pricing as a host? Why or why not?",
            "Connect to regression: What is the target variable? (Revenue? Bookings? Occupancy?)",
            "Ties nicely to the pandemic disruption theme from Zillow"
        ]
    },

    # WEEK 2A: CAPITAL ONE
    {
        "week": 2,
        "topic": "Tree-Based Predictions",
        "id": "Week2A_CapitalOne",
        "title": "Capital One: From Credit Cards to AI Pioneer",
        "content_paragraphs": [
            "Capital One was founded in 1988 on a radical premise: use data and analytics to make better credit decisions than competitors. Today, machine learning is integrated into 'almost every facet of the business,' from fraud detection to credit risk assessment to anti-money laundering compliance.",

            "The company's approach to model selection reveals a critical tension in financial AI. When Capital One built its anti-money laundering (AML) system, they initially investigated logistic regression, XGBoost, and recurrent neural networks. They chose a random forest model because it represented 'the best balance of model accuracy, performance, and explainability for this use case.'",

            '"Being in a heavily regulated environment, we want to make sure that we are not just meeting the regulatory requirements, but that we help set the standard for what fair and ethical machine learning deployment looks like," explained a Capital One executive. The choice was not just about accuracy-it was about being able to explain decisions to regulators and customers.',

            "This explainability requirement shapes which models Capital One can deploy. Research has shown that XGBoost achieves 99.88% accuracy in fraud detection with an AUC-ROC score of 1.0, but that performance comes with a transparency cost. By investing heavily in documentation and using open-source tools that enable others to train identical models and verify results, Capital One created systems where 'stakeholders-such as federal regulators-can see and understand the work.'",

            "The regulatory constraints are real. The Equal Credit Opportunity Act (ECOA) and Fair Credit Reporting Act (FCRA) require that when you deny someone credit, you must explain why. 'As the models and algorithms become more advanced, it becomes harder to explain to customers and regulators alike how the underlying models work, and with complexity also comes an increased risk of harm from biased, untested, or unfair outcomes.'",

            "Capital One now explores synthetic data generation and federated learning to enhance privacy, tabular data solutions for fraud detection, explainability methods including topical data analysis and model introspection, and sequence modeling for credit risk. The constraint of explainability has become a driver of innovation.",
        ],
        "key_facts": [
            ("Founded", "1988 (data-first from day one)"),
            ("AML model choice", "Random Forest (over XGBoost, RNN)"),
            ("Reason for choice", "'Best balance of accuracy, performance, and explainability'"),
            ("Regulatory requirements", "ECOA, FCRA require explainable decisions"),
            ("XGBoost fraud accuracy", "99.88% (but less explainable)"),
            ("Research focus", "Synthetic data, federated learning, explainability"),
        ],
        "questions": [
            {
                "question": "Capital One chose Random Forest over XGBoost for AML despite XGBoost's superior accuracy. When is this trade-off worth making?",
                "notes": [
                    "Regulated industries require explainability",
                    "Trees have clear decision paths that can be audited",
                    "XGBoost is powerful but harder to explain",
                    "Trade-off: Some accuracy for interpretability"
                ]
            },
            {
                "question": "ECOA and FCRA require that credit denials be explained. How does this constraint shape which models can be deployed?",
                "notes": [
                    "Trees: 'You were denied because income < $X and debt > $Y'",
                    "Neural networks: Black box-hard to explain specific decisions",
                    "Regulation effectively mandates certain model architectures",
                    "Innovation must happen within constraints"
                ]
            },
            {
                "question": "ML models can inadvertently encode historical biases. How should a company balance predictive accuracy with fairness?",
                "notes": [
                    "Historical data reflects historical discrimination",
                    "Proxies for protected classes (zip code to race)",
                    "Fairness metrics: demographic parity, equalized odds",
                    "Tension: Removing predictive features may reduce accuracy"
                ]
            },
            {
                "question": "Capital One uses synthetic data and federated learning for privacy. Why might these matter for financial ML?",
                "notes": [
                    "Real customer data is sensitive",
                    "Synthetic data enables testing without privacy risk",
                    "Federated learning keeps data distributed",
                    "Regulatory pressure on data handling increasing"
                ]
            }
        ],
        "key_insight": "Tree models remain industry workhorses because they balance accuracy with interpretability. In regulated industries, being able to explain your decision is as important as making the right one.",
        "primary_source": "d3.harvard.edu/platform-rctom/submission/machine-learning-in-credit-assessment-at-capital-one/",
        "teaching_tips": [
            "Great case for discussing the accuracy-interpretability trade-off",
            "Ask: Would you want a black-box model deciding your loan application?",
            "Connect to lecture: Feature importance in trees equals explainability",
            "Fairness discussion can be extended or shortened based on time"
        ]
    },

    # WEEK 2B: JPMORGAN
    {
        "week": 2,
        "topic": "Tree-Based Predictions",
        "id": "Week2B_JPMorgan",
        "title": "JPMorgan COIN: Contract Intelligence",
        "content_paragraphs": [
            "Before COIN, JPMorgan faced a crushing operational burden: lawyers and loan officers spent approximately 360,000 hours annually reviewing commercial loan agreements. That is equivalent to more than 41 years of nonstop work-all devoted to a 'mind-numbing job' of interpreting complex legal documents.",

            "JPMorgan's COIN (Contract Intelligence) system changed that equation. Using natural language processing and machine learning, the system 'reviews documents in seconds, is less error-prone and never asks for vacation.' In its early deployment, COIN derived '150 attributes from 12,000 commercial credit agreements and contracts in only a few seconds.'",

            "The system employs unsupervised learning to identify and categorize repeated clauses in credit contracts. Beyond speed, it actually improved quality: COIN 'has helped JPMorgan cut down on loan-servicing mistakes, most of which stemmed from human error in interpreting 12,000 new wholesale contracts per year.' The algorithm proved more accurate than human lawyers, making the investment about quality as much as cost.",

            "The operational impact was immediate: work that consumed 360,000 hours annually now takes seconds. But this raises profound questions about the legal profession. What happened to the lawyers who previously did this work? In the short term, JPMorgan intended to deploy COIN for more complex filings-credit-default swaps and custody agreements. In the medium term, the bank hopes to use machine learning to interpret altogether new regulations.",

            "The liability question looms large. When AI misses a critical contract term or misinterprets an obligation, who bears responsibility? Currently, lawyers still sign off on AI-reviewed documents, retaining liability. But as AI's role expands, the legal framework remains unclear.",

            "JPMorgan's success with COIN reflects broader patterns in professional services automation. Document-heavy, pattern-recognizable tasks-due diligence, compliance review, patent search-are increasingly tractable for NLP. Tasks requiring judgment, relationship management, and courtroom advocacy remain human domains. For now.",
        ],
        "key_facts": [
            ("Previous manual effort", "360,000 hours annually"),
            ("Current processing time", "Seconds"),
            ("Early results", "150 attributes from 12,000 contracts in seconds"),
            ("Quality improvement", "Reduced loan-servicing mistakes from human error"),
            ("Model approach", "Unsupervised learning for clause categorization"),
            ("Future applications", "Credit-default swaps, custody agreements, new regulations"),
        ],
        "questions": [
            {
                "question": "COIN saved 360,000 hours of lawyer time-equivalent to 41 years of continuous work. What happened to those lawyers and their roles?",
                "notes": [
                    "Shifted to higher-value work (negotiation, strategy)",
                    "More deals processed with same headcount",
                    "Junior lawyer roles most affected (entry-level review work)",
                    "New roles: AI training, validation, exception handling"
                ]
            },
            {
                "question": "COIN uses unsupervised learning to categorize contract clauses. Why unsupervised rather than supervised learning?",
                "notes": [
                    "No pre-existing labels for all clause types",
                    "Discovers patterns humans might not think to look for",
                    "Contracts evolve-new clause types emerge",
                    "Supervised would require massive labeling effort"
                ]
            },
            {
                "question": "When AI misses a critical contract term, who bears liability-JPMorgan, the AI vendor, or the lawyer who approved it?",
                "notes": [
                    "Currently: Lawyers still sign off (retain liability)",
                    "Future: Unclear as AI role expands",
                    "Insurance implications",
                    "Regulatory guidance evolving"
                ]
            },
            {
                "question": "Which other professional services might be transformed by NLP? Which are resistant?",
                "notes": [
                    "Transformable: Due diligence, compliance, auditing, medical records",
                    "Resistant: Courtroom advocacy, client relationships, negotiation",
                    "Key: Document-heavy, pattern-recognizable tasks are tractable",
                    "Judgment and relationships remain human domains"
                ]
            }
        ],
        "key_insight": "NLP can automate document-heavy professional work, but stakes are high when mistakes have legal consequences. Human oversight remains essential.",
        "primary_source": "abajournal.com/news/article/jpmorgan_chase_uses_tech_to_save_360000_hours_of_annual_work_by_lawyers_and",
        "teaching_tips": [
            "Good for discussing professional services automation",
            "Connects to broader workforce transformation discussion",
            "Connect to lecture: NLP for document understanding",
            "Ask: What would you want AI to review in a contract you are signing?"
        ]
    },

    # WEEK 2C: HEALTHCARE FDA
    {
        "week": 2,
        "topic": "Tree-Based Predictions",
        "id": "Week2C_Healthcare",
        "title": "Healthcare AI: FDA-Approved Medical Devices",
        "content_paragraphs": [
            "By mid-2025, the FDA had authorized over 1,200 AI/ML-enabled medical devices-a number growing rapidly as healthcare embraces artificial intelligence. The vast majority, roughly 75-80%, reside in radiology, with cardiology accounting for about 10% and neurology, hematology, and other specialties making up the rest.",

            "The regulatory pathway matters enormously. Overall, 97% of these devices were cleared via the 510(k) pathway-a mechanism that streamlines market entry by demonstrating 'substantial equivalence' to a predicate device already on the market. This pathway 'does not require independent clinical data demonstrating performance or safety.'",

            "The validation gap is striking. Of radiology devices with submission documentation, only 5% underwent prospective testing, 8% included a human-in-the-loop evaluation, and 29% incorporated any clinical testing. The 510(k) process caps the number of indications for which applicants can seek approval at a given time, potentially limiting thorough validation.",

            "Viz.ai exemplifies both the promise and the regulatory evolution. The company created software that analyzes CT images to detect potential strokes and automatically notify neurovascular specialists. The FDA cleared it through the De Novo pathway-for devices with no existing predicate-creating a new regulatory classification that subsequent computer-aided triage software can use.",

            "The clinical impact has been substantial. Viz.ai is now deployed in over 1,400 hospitals, covering more than 220 million lives. Studies show 66-minute faster treatment times with the platform-critical when brain tissue dies rapidly during strokes. The company now has 13 FDA-cleared algorithms for stroke and neurocritical care.",

            "But questions remain about the broader ecosystem. While AI promises faster, more consistent diagnosis, the lack of prospective testing for most devices raises safety concerns. The 510(k) pathway enables rapid innovation but may not catch all failure modes before deployment. As one analysis noted: 'While efficient, the 510(k) pathway does not require independent clinical data demonstrating performance or safety.'",
        ],
        "key_facts": [
            ("FDA-authorized AI devices", "1,200+ by mid-2025"),
            ("Radiology concentration", "75-80% of all devices"),
            ("510(k) pathway usage", "97% of approvals"),
            ("Prospective testing", "Only 5% of radiology devices"),
            ("Viz.ai hospitals", "1,400+ (220 million lives covered)"),
            ("Stroke treatment improvement", "66 minutes faster"),
        ],
        "questions": [
            {
                "question": "Why has radiology become the dominant domain for FDA-approved AI (75-80% of devices)?",
                "notes": [
                    "Clear ground truth (biopsy, follow-up imaging)",
                    "Abundant digital data (images already stored)",
                    "Defined tasks (detect tumor, measure size)",
                    "Workflow integration is straightforward"
                ]
            },
            {
                "question": "Only 5% of approved radiology devices underwent prospective testing. What are the risks of this approach?",
                "notes": [
                    "510(k): 'Substantially equivalent' to existing device",
                    "Training data may not match deployment population",
                    "Edge cases and rare conditions may not be covered",
                    "Post-market surveillance becomes critical"
                ]
            },
            {
                "question": "Viz.ai reduces stroke treatment time by 66 minutes. How would you calculate the health and economic value of this improvement?",
                "notes": [
                    "Brain tissue dies rapidly during stroke",
                    "Better outcomes reduce long-term care costs",
                    "Quality-adjusted life years (QALYs) saved",
                    "Compare device cost to treatment cost savings"
                ]
            },
            {
                "question": "What is the appropriate role for AI in medical diagnosis-replacement, augmentation, or second opinion?",
                "notes": [
                    "Current FDA stance: AI as 'decision support'",
                    "Liability questions: Who is responsible for errors?",
                    "Radiologist + AI often outperforms either alone",
                    "Risk tolerance varies by condition severity"
                ]
            }
        ],
        "key_insight": "Classification models in healthcare must balance speed-to-market with safety validation. The 510(k) pathway enables innovation but may not catch all failure modes.",
        "primary_source": "jamanetwork.com/journals/jamanetworkopen/fullarticle/2841066",
        "teaching_tips": [
            "Great case for discussing classification metrics in high-stakes settings",
            "Ask: What false positive rate is acceptable for cancer screening?",
            "Connect to lecture: Precision vs. recall trade-off",
            "Can extend to discuss algorithmic bias in healthcare"
        ]
    },

    # WEEK 3A: SPOTIFY/NETFLIX
    {
        "week": 3,
        "topic": "Clustering & Collaborative Filtering",
        "id": "Week3A_Spotify",
        "title": "Spotify & Netflix: Taste Communities",
        "content_paragraphs": [
            "Traditional market research segments customers by demographics-age, income, location, education. Spotify and Netflix discovered that these variables are poor predictors of taste. A teenager in Brazil and a retiree in Japan might have more in common than two neighbors of the same age.",

            "Spotify pioneered the application of collaborative filtering in music-what they call the 'Netflix approach.' The basic principle: 'We can understand songs to recommend to a user by looking at what other users with similar tastes are listening to.' If user A has enjoyed songs X, Y, and Z, and user B has enjoyed songs X and Y but has not heard song Z, the system recommends song Z to user B.",

            "The scale of Spotify's collaborative filtering is remarkable. The model is trained on a sample of 700 million user-generated playlists, selected from the 9 billion total playlists on the platform. As one analysis noted, this data represents 'the passion, care, love, and time users spend creating playlists.' By maintaining a massive user-item interaction matrix covering all users and tracks, Spotify can identify which songs are similar (listened to by similar users) and which users are similar (listen to the same songs).",

            "Rather than using traditional genre labels, Spotify employs 'Taste Analysis Data' to establish Taste Profiles. This technology groups music into clusters based on listening behavior, not human categorization. Through this analysis, Spotify builds 'a map of music, allowing songs to form clusters based on user behaviors'-revealing relationships that genre labels might miss.",

            "Netflix takes a similar approach, identifying over 1,300 'taste communities' that cut across geography and demographics. A viewer's cluster membership predicts their preferences far better than knowing their age or location. This insight has profound implications for content creation: global content can find global audiences, and niche programming becomes viable through the long tail.",

            "The methods work because behavior reveals preferences directly, while demographics are mere proxies. Two 25-year-olds may have completely different tastes. But two people who both loved the same obscure documentary are likely to enjoy similar content-regardless of where they live or how old they are.",
        ],
        "key_facts": [
            ("Spotify training data", "700 million user-generated playlists"),
            ("Total Spotify playlists", "9 billion"),
            ("Netflix taste communities", "1,300+"),
            ("Key insight", "Behavior predicts taste better than demographics"),
            ("Approach", "Collaborative filtering (the 'Netflix approach')"),
            ("Clustering basis", "Listening/viewing behavior, not genre labels"),
        ],
        "questions": [
            {
                "question": "Spotify trains on 700 million playlists from 9 billion total. What makes user-generated playlists valuable for understanding taste?",
                "notes": [
                    "Playlists represent intentional curation",
                    "Reveal which songs 'go together' in users' minds",
                    "More signal than passive listening",
                    "Capture context (workout, study, party)"
                ]
            },
            {
                "question": "Netflix found that a teenager in Brazil and a retiree in Japan might be in the same 'taste cluster.' What are the implications for content creation?",
                "notes": [
                    "Global content can find global audiences",
                    "Don't need 'content for seniors' or 'content for teens'",
                    "Niche content becomes viable (long tail)",
                    "Reduces reliance on demographic assumptions"
                ]
            },
            {
                "question": "Spotify uses behavior-based clusters instead of genre labels. Why might behavior be more useful than human-defined categories?",
                "notes": [
                    "Genres are subjective and overlapping",
                    "Behavior captures implicit preferences",
                    "Cross-genre relationships emerge naturally",
                    "Avoids pigeonholing by human categorization"
                ]
            },
            {
                "question": "What are the limitations of collaborative filtering? When might it fail?",
                "notes": [
                    "Cold start problem: new users have no history",
                    "New items have no interactions",
                    "Filter bubbles: only recommends similar content",
                    "Popularity bias: popular items dominate"
                ]
            }
        ],
        "key_insight": "Unsupervised learning reveals structure humans would not think to look for. Behavior-based segments often outperform demographic assumptions.",
        "primary_source": "music-tomorrow.com/blog/how-spotify-recommendation-system-works-complete-guide",
        "teaching_tips": [
            "Ask students what demographic segment they are in vs. their actual Netflix taste",
            "Good discussion on cold start problem",
            "Connect to lecture: Matrix factorization, collaborative filtering",
            "Can discuss filter bubbles as a downside"
        ]
    },

    # WEEK 4A: UBER
    {
        "week": 4,
        "topic": "Reinforcement Learning",
        "id": "Week4A_Uber",
        "title": "Uber: Marketplace Optimization at Scale",
        "content_paragraphs": [
            "Uber operates one of the world's largest real-time marketplaces, making over 1 million pricing decisions per second with sub-50ms latency requirements, serving millions of drivers across 70+ countries. At this scale, the difference between good and optimal matching translates to billions of dollars.",

            "The naive approach would match each rider to the nearest available driver-a greedy algorithm. But this creates problems. If you send all drivers to the airport because demand spikes there, downtown gets starved. A greedy algorithm 'without an understanding of subsequent likely outcomes might create balance at the time of the match, but may cause imbalances in other parts of the city in the future, leading to longer wait times or surge pricing elsewhere.'",

            "Uber's solution uses reinforcement learning-specifically a DQN (Deep Q-Network) inspired approach-to learn value functions and optimize driver matching. The system models matching as a Markov Decision Process (MDP) where 'the agent takes collective decisions to match drivers to riders in a particular order.' This is now deployed in over 400 cities globally-'the largest production deployment of a reinforcement learning algorithm for matching in the ridesharing marketplace.'",

            "The RL system does not just optimize immediate matches; it anticipates future demand. It might send a driver to pick up a slightly farther rider to maintain balance across the city. The goal is marketplace-level optimization, not just individual ride optimization.",

            "Uber's systems also use reinforcement learning for driver incentives, evaluating how drivers respond to different incentive structures-quests, guarantees, targeted bonuses-and continuously adjusting them for optimal outcomes. The company built a simulation platform hosting 'a simulated world with driver-partners and riders, mimicking scenarios in the real world,' allowing engineers to rapidly prototype and test new features in a risk-free environment.",

            "Dynamic pricing-surge-remains controversial but serves an economic purpose. Higher prices reduce demand (fewer riders request rides) while increasing supply (more drivers come online). Without surge pricing, demand exceeds supply, and no cars are available at all. The RL system must balance multiple objectives: rider wait times, driver earnings, marketplace efficiency, and customer experience.",
        ],
        "key_facts": [
            ("Scale", "1M+ pricing decisions per second"),
            ("Latency requirement", "Sub-50ms"),
            ("RL approach", "Deep Q-Network (DQN) inspired"),
            ("Deployment", "400+ cities (largest RL matching deployment)"),
            ("Framework", "Markov Decision Process (MDP)"),
            ("Testing", "Simulation platform with virtual drivers/riders"),
        ],
        "questions": [
            {
                "question": "A greedy algorithm would match each rider to the nearest driver. Why is this suboptimal for the overall marketplace?",
                "notes": [
                    "Creates supply deserts in some areas",
                    "All drivers rush to high-demand areas (airport)",
                    "Doesn't consider future demand",
                    "Local optimization hurts global outcomes"
                ]
            },
            {
                "question": "Uber uses RL to balance supply and demand across a city. What happens if RL optimizes too aggressively for one metric?",
                "notes": [
                    "Optimize only wait time: Drivers hate it (short trips)",
                    "Optimize only driver earnings: Riders wait too long",
                    "Need multi-objective optimization",
                    "Goodhart's Law: When a measure becomes a target..."
                ]
            },
            {
                "question": "Surge pricing is controversial but Uber argues it is necessary. How would you explain the economic rationale?",
                "notes": [
                    "Higher prices reduce demand (fewer riders)",
                    "Higher prices increase supply (more drivers online)",
                    "Without surge: demand > supply = no cars available",
                    "Surge is allocation mechanism, not just profit"
                ]
            },
            {
                "question": "Uber built a simulation platform to test matching algorithms. Why simulate rather than test in production?",
                "notes": [
                    "Risk-free experimentation",
                    "Can test rare scenarios (major events, weather)",
                    "Faster iteration than real-world A/B tests",
                    "Protects customer experience during testing"
                ]
            }
        ],
        "key_insight": "RL shines in sequential decisions where actions affect future states. Greedy optimization often fails when local decisions have global consequences.",
        "primary_source": "uber.com/blog/reinforcement-learning-for-modeling-marketplace-balance/",
        "teaching_tips": [
            "Most students use rideshare-ask about surge pricing experiences",
            "Good for explaining the explore/exploit trade-off",
            "Connect to lecture: States, actions, rewards framework",
            "Can discuss ethical concerns about algorithmic pricing"
        ]
    },

    # WEEK 4B: ALPHAGO
    {
        "week": 4,
        "topic": "Reinforcement Learning",
        "id": "Week4B_AlphaGo",
        "title": "AlphaGo: The Game That Changed AI",
        "content_paragraphs": [
            "In March 2016, AlphaGo faced Lee Sedol-a 9-dan professional Go player and one of the strongest players in history-in a five-game match watched by over 200 million people. Go has more possible board positions than atoms in the universe, making brute-force search impossible. AlphaGo won 4-1, marking a watershed moment for artificial intelligence.",

            "The match is often compared to Deep Blue versus Garry Kasparov in 1997, but there is a crucial difference. Chess can be solved through search-evaluating millions of positions per second. Go cannot. AlphaGo required a fundamentally different approach: using 'value networks' to evaluate board positions and 'policy networks' to select moves, trained through 'a novel combination of supervised learning from human expert games, and reinforcement learning from games of self-play.'",

            "Move 37 in the second game became legendary. AlphaGo played a shoulder hit at the fifth line-a move so unconventional that AlphaGo itself calculated only a 1:10,000 probability that a human would make it. Commentators were initially baffled, then described it as 'beautiful' and 'creative.' Lee Sedol needed 15 minutes to contemplate his response. The strategy 'was not taken out of a database of publicly known moves. Move 37 was new to the 5,500-year history of Go.'",

            "But AlphaGo's initial training relied on human games. DeepMind then created AlphaGo Zero, which 'starts off with a neural network that knows nothing about the game of Go' and learns 'purely from self-play.' The results were stunning: 'After just three days of self-play training, AlphaGo Zero emphatically defeated the previously published version of AlphaGo-which had itself defeated 18-time world champion Lee Sedol-by 100 games to 0.'",

            "After 40 days of self-training, AlphaGo Zero outperformed even the 'Master' version that had beaten every top professional. David Silver and colleagues showed that 'starting from random moves, it can reach superhuman level in just a couple of days of training and five million games of self-play.' The system discovered strategies humans never considered-validated by Move 37's success.",

            "The implications extend far beyond games. As one analysis noted: Move 37 'highlighted how AI could transcend traditional human approaches in complex decision-making environments... reshaping our understanding of AI's potential in strategic reasoning and its application beyond board games.' The techniques have since been applied to protein folding, chip design, and other domains where simulation is cheap and objectives are clear.",
        ],
        "key_facts": [
            ("Match result", "AlphaGo 4, Lee Sedol 1 (March 2016)"),
            ("Viewers", "200+ million"),
            ("Move 37 probability", "1:10,000 chance a human would play it"),
            ("AlphaGo Zero vs. original", "100-0 after 3 days of self-play"),
            ("Training for superhuman level", "5 million games of self-play"),
            ("Lee Sedol's rank", "9-dan (one of history's strongest players)"),
        ],
        "questions": [
            {
                "question": "Go has more possible positions than atoms in the universe. Why did reinforcement learning succeed where brute-force search could not?",
                "notes": [
                    "Cannot enumerate all positions (unlike chess)",
                    "RL learns value function: Which positions are good?",
                    "Does not need to search everything-just promising paths",
                    "Neural network generalizes from seen positions to unseen"
                ]
            },
            {
                "question": "Move 37 was described as 'beautiful' and 'creative,' yet AlphaGo has no consciousness. Can machines be creative?",
                "notes": [
                    "Move 37 was genuinely novel-not in any database",
                    "Self-play discovers strategies humans never considered",
                    "Creativity might be exploration of possibility space",
                    "Human creativity also builds on patterns"
                ]
            },
            {
                "question": "AlphaGo Zero learned without human data and beat the human-trained version 100-0. What does this suggest about human expertise?",
                "notes": [
                    "Humans may have suboptimal conventional wisdom",
                    "Self-play can discover superior strategies",
                    "But: Humans learn with far less data",
                    "Human expertise is sample-efficient; AI is not"
                ]
            },
            {
                "question": "AlphaGo's techniques have been applied to protein folding and chip design. What makes a problem suitable for this approach?",
                "notes": [
                    "Clear objective function (win/lose, structure fits)",
                    "Ability to simulate/evaluate outcomes cheaply",
                    "Large but structured search space",
                    "Not suitable: Open-ended problems, unclear goals"
                ]
            }
        ],
        "key_insight": "RL can discover strategies that surpass human expertise when rules are clear and simulation is cheap. Self-play enables learning without human bias.",
        "primary_source": "en.wikipedia.org/wiki/AlphaGo_versus_Lee_Sedol",
        "teaching_tips": [
            "Show the Move 37 video clip if time permits",
            "AlphaGo Zero is the more important result pedagogically",
            "Connect to lecture: Self-play as RL without human labels",
            "Good transition to discussing AI capabilities and limitations"
        ]
    },

    # WEEK 5A: ALPHAFOLD
    {
        "week": 5,
        "topic": "Neural Networks",
        "id": "Week5A_AlphaFold",
        "title": "AlphaFold: Solving Biology's 50-Year Grand Challenge",
        "content_paragraphs": [
            "For 50 years, scientists struggled with the 'protein folding problem'-predicting how a protein's amino acid sequence folds into its three-dimensional structure. This structure determines function, making the problem central to understanding disease and designing drugs. Progress was painfully slow.",

            "In November 2020, DeepMind's AlphaFold 2 entered the CASP14 competition-the biennial Olympics of protein structure prediction. The results shocked the scientific community. AlphaFold made the best prediction for 88 out of 97 targets, achieving a median Global Distance Test (GDT) score of 92.4 out of 100. More remarkably, it predicted structures down to atomic accuracy with median error (RMSD_95) of less than 1 Angstrom-three times more accurate than the next best system and comparable to experimental methods.",

            'Nobel Prize winner and structural biologist Venki Ramakrishnan called the result "a stunning advance on the protein folding problem," adding that "It has occurred decades before many people in the field would have predicted."',

            "In 2022, DeepMind released predictions for more than 200 million protein structures-'nearly all catalogued proteins known to science.' The AlphaFold Protein Structure Database makes this data freely available. So far, it has attracted over 3 million users from over 190 countries.",

            "The impact has been profound. AlphaFold is being used to tackle problems including antimicrobial resistance, crop resilience, and heart disease. Over 30% of AlphaFold-related research focuses on understanding disease. The work's scientific and societal value was recognized in 2024 when Demis Hassabis and John Jumper of DeepMind shared half of the Nobel Prize in Chemistry 'for protein structure prediction.'",

            "DeepMind's decision to freely release 200+ million structures raises strategic questions. The company could have monetized this breakthrough directly. Instead, they chose to accelerate scientific progress and demonstrate capabilities-drawing talent, generating goodwill, and showcasing what AI can achieve. For drug companies, AlphaFold has leveled the playing field: all have access, so competitive advantage now lies in integration speed and what comes after structure prediction.",
        ],
        "key_facts": [
            ("Problem duration", "50 years unsolved"),
            ("CASP14 score", "92.4/100 GDT (3x more accurate than next best)"),
            ("Atomic accuracy", "< 1 Angstrom median error"),
            ("Structures released", "200+ million (free)"),
            ("Database users", "3+ million from 190+ countries"),
            ("Nobel Prize", "2024 Chemistry (Hassabis and Jumper)"),
        ],
        "questions": [
            {
                "question": "AlphaFold solved a problem biologists worked on for 50 years. What made it tractable for neural networks?",
                "notes": [
                    "Large dataset: Protein Data Bank had 180,000+ structures",
                    "Clear objective: Predict atomic coordinates",
                    "Physical constraints: Chemistry rules narrow possibilities",
                    "Attention mechanism captures long-range dependencies"
                ]
            },
            {
                "question": "DeepMind made AlphaFold freely available instead of monetizing it. What are the strategic implications?",
                "notes": [
                    "PR and prestige value (talent attraction)",
                    "Accelerates scientific progress (goodwill)",
                    "AlphaFold itself is not a product-it is a capability demonstration",
                    "Google/Alphabet values long-term research over short-term revenue"
                ]
            },
            {
                "question": "AlphaFold predicts structures but does not explain why proteins fold as they do. Is this 'understanding'? Does it matter?",
                "notes": [
                    "Prediction vs. explanation debate",
                    "Practical utility does not require understanding",
                    "Scientists still want to know mechanisms",
                    "Similar debate in all of ML"
                ]
            },
            {
                "question": "All drug companies now have access to AlphaFold. How does this change pharmaceutical competitive dynamics?",
                "notes": [
                    "Levels playing field (all have access)",
                    "Advantage goes to those who integrate fastest",
                    "Shifts bottleneck to other parts of drug development",
                    "Still need wet lab validation"
                ]
            }
        ],
        "key_insight": "Neural networks excel with abundant data and clear objectives-even for problems humans could not solve. Open-sourcing can accelerate science while demonstrating capabilities.",
        "primary_source": "deepmind.google/discover/blog/alphafold-a-solution-to-a-50-year-old-grand-challenge-in-biology/",
        "teaching_tips": [
            "Nobel Prize 2024 in Chemistry went to AlphaFold-very current",
            "Show visualizations of protein structures if possible",
            "Good for discussing the 'black box' vs. understanding debate",
            "Connect to lecture: Attention mechanism (foreshadows Week 7)"
        ]
    },

    # WEEK 5B: DEEPMIND DATA CENTERS
    {
        "week": 5,
        "topic": "Neural Networks",
        "id": "Week5B_DeepMind",
        "title": "DeepMind: 40% Reduction in Data Center Cooling Costs",
        "content_paragraphs": [
            "Google's data centers consume massive amounts of energy, with cooling often accounting for 30-40% of total power usage in less efficient facilities. In 2016, DeepMind applied its machine learning expertise to a seemingly mundane problem: optimizing cooling systems. The results were dramatic.",

            "DeepMind's system achieved a 40% reduction in the amount of energy used for cooling, which translated to a 15% reduction in overall Power Usage Effectiveness (PUE)-the ratio of total building energy to IT energy. This marked the lowest PUE ever recorded at the site.",

            "The technical approach involved training an ensemble of deep neural networks on historical data from thousands of sensors within the data center-temperatures, power consumption, pump speeds, setpoints, and more. The networks were trained to predict future PUE, temperature, and pressure over the next hour. After implementation, the team achieved PUE predictions with only 0.4% error, at a value PUE of 1.1 (remarkably close to the theoretical minimum of 1.0).",

            "The system takes sensor readings and recommends actions every five minutes, with humans reviewing before implementation. This human-in-the-loop approach addresses a crucial concern: cooling failures in data centers can cause equipment damage worth millions. The AI makes recommendations; humans verify safety before execution.",

            "Eighteen months after initial development, the models were piloted at multiple facilities, consistently delivering around 30% average cooling energy savings. At Google's scale, this translates to millions of dollars saved per year in each large data center. As one DeepMind researcher noted: 'Every improvement in data center efficiency reduces total emissions into our environment.'",

            "The success led to broader applications. DeepMind partnered with the UK National Grid to explore AI for energy grid optimization. But applying AI to critical infrastructure raises the stakes: grid failures cause blackouts affecting millions, requiring longer validation periods and more robust safety margins than data center cooling. The technology transfers, but the risk tolerance doesn't.",
        ],
        "key_facts": [
            ("Cooling energy reduction", "40%"),
            ("Overall PUE improvement", "15%"),
            ("PUE prediction accuracy", "0.4% error"),
            ("Data inputs", "Thousands of sensors"),
            ("Recommendation frequency", "Every 5 minutes"),
            ("Human oversight", "All recommendations reviewed before execution"),
            ("Broader application", "UK National Grid partnership"),
        ],
        "questions": [
            {
                "question": "Data center cooling seems like an engineering problem. Why did neural networks outperform traditional optimization?",
                "notes": [
                    "Thousands of sensors = high-dimensional input",
                    "Complex, non-linear interactions between variables",
                    "Traditional control theory struggles with this complexity",
                    "Neural networks can learn from operational data"
                ]
            },
            {
                "question": "The system makes recommendations every 5 minutes but humans review before execution. When is human oversight essential?",
                "notes": [
                    "Safety-critical systems (cooling failure = equipment damage)",
                    "Novel situations outside training data",
                    "Building trust before full automation",
                    "Regulatory or liability requirements"
                ]
            },
            {
                "question": "Google saved 40% on cooling. What other industrial processes might benefit from similar approaches?",
                "notes": [
                    "Manufacturing (process optimization)",
                    "HVAC in large buildings",
                    "Chemical plants (yield optimization)",
                    "Key requirement: sensors + clear objective"
                ]
            },
            {
                "question": "DeepMind partnered with UK National Grid. What is different about AI for critical infrastructure vs. data centers?",
                "notes": [
                    "Failure is not an option (blackouts)",
                    "Regulatory constraints and public accountability",
                    "Longer validation periods required",
                    "Technology transfers, but risk tolerance doesn't"
                ]
            }
        ],
        "key_insight": "Neural networks can optimize complex systems with many interacting variables. Human oversight remains important for safety-critical applications.",
        "primary_source": "deepmind.google/blog/deepmind-ai-reduces-google-data-centre-cooling-bill-by-40/",
        "teaching_tips": [
            "Good case for industrial AI applications",
            "Emphasize the human-in-the-loop approach",
            "Connect to lecture: Neural networks as function approximators",
            "Can discuss energy/climate implications"
        ]
    },

    # WEEK 6A: TESLA
    {
        "week": 6,
        "topic": "Convolutional Neural Networks",
        "id": "Week6A_Tesla",
        "title": "Tesla Autopilot: The Vision-Only Bet",
        "content_paragraphs": [
            "Tesla made a bold bet: autonomous driving can be achieved with cameras alone-no LIDAR, no radar. In 2021, they removed radar entirely, transitioning to a vision-only system. This approach diverges sharply from competitors like Waymo, who use LIDAR, cameras, and radar together for redundant perception.",

            "Tesla's technical architecture centers on HydraNet-a single massive neural network that processes inputs from 8 cameras simultaneously. The company coined the term for its 'large camera perception network that has many heads solving different tasks.' There is a shared backbone that processes raw camera images into features, and then multiple output heads that produce specific predictions-traffic light detection, lane prediction, obstacle recognition, and more.",

            "The data advantage is Tesla's key differentiator. The company processes 400,000 video clips per second from its global fleet. 'Thousands of people are driving in Tesla everyday! The data is coming from the fleet.' Every piece of data is collected, labeled, and used for training. Tesla's 'shadow mode' runs neural networks in the background, comparing their decisions to human drivers and flagging disagreements for training. When drivers intervene or take over from Autopilot, these moments become valuable training examples.",

            "The training infrastructure is massive. Tesla's custom Dojo supercomputer features D1 chips with 362 teraflops each. Their primary data center houses 14,000 GPUs, with plans to expand to 50,000 NVIDIA H100 GPUs in the Cortex cluster. This investment reflects the company's belief that software-trained on vast amounts of real-world driving data-will eventually outperform sensor-based redundancy.",

            "FSD v12's launch in late 2023 represented a paradigm shift. The approximately 300,000 lines of control code collapsed to roughly 2,000-3,000 lines needed simply to activate and manage the neural networks. The system moved from rule-based to end-to-end learning.",

            "But the approach remains controversial. Tesla's 'Full Self-Driving' name suggests Level 4 or 5 autonomy, but the system is SAE Level 2-drivers must remain attentive. Critics argue that cameras struggle in darkness and adverse weather where LIDAR excels. Tesla's counterargument: humans drive with vision alone, so cameras should suffice. The debate is unresolved.",
        ],
        "key_facts": [
            ("Sensor approach", "8 cameras only (no LIDAR, no radar since 2021)"),
            ("Architecture", "HydraNet (multi-task neural network)"),
            ("Fleet data processing", "400,000 video clips per second"),
            ("Training infrastructure", "14,000 GPUs, expanding to 50,000 H100s"),
            ("FSD v12 code reduction", "~300,000 lines to ~2,000-3,000 lines"),
            ("Autonomy level", "SAE Level 2 (driver must remain attentive)"),
        ],
        "questions": [
            {
                "question": "Tesla bet on cameras while competitors use LIDAR. What are the trade-offs of each approach?",
                "notes": [
                    "Cameras: Cheap, scalable, but struggle in darkness/weather",
                    "LIDAR: Precise depth, works at night, but expensive",
                    "Tesla's bet: Neural networks can extract depth from video",
                    "Competitors' bet: Redundant sensors are safer"
                ]
            },
            {
                "question": "Tesla trains on data from 400,000 video clips per second from its fleet. What advantages and risks does this create?",
                "notes": [
                    "Advantage: Millions of cars = massive data",
                    "Advantage: Real-world edge cases (not just test tracks)",
                    "Risk: Privacy concerns (always recording)",
                    "Risk: Selection bias (where Tesla owners drive)"
                ]
            },
            {
                "question": "FSD v12 collapsed 300,000 lines of code to ~3,000 lines by using end-to-end learning. What are the implications?",
                "notes": [
                    "Less hand-coded rules = less human bias/error",
                    "But: Harder to understand why system acts",
                    "Debugging neural networks is harder than debugging code",
                    "Shift from engineering to training"
                ]
            },
            {
                "question": "Tesla's 'Full Self-Driving' is actually Level 2. How should companies communicate AI capabilities and limitations?",
                "notes": [
                    "FSD name suggests Level 4 or 5",
                    "Marketing vs. safety communication tension",
                    "Regulatory scrutiny increasing",
                    "Over-promising erodes trust"
                ]
            }
        ],
        "key_insight": "Computer vision enables machines to perceive the world, but perception is not understanding. Different companies make fundamentally different bets on the path to autonomy.",
        "primary_source": "thinkautonomous.ai/blog/how-tesla-autopilot-works/",
        "teaching_tips": [
            "Very current and polarizing-students will have opinions",
            "Good for discussing sensor fusion vs. single-modality approaches",
            "Connect to lecture: CNNs processing camera images",
            "Can discuss the marketing vs. capability gap"
        ]
    },

    # WEEK 6B: WAYMO
    {
        "week": 6,
        "topic": "Convolutional Neural Networks",
        "id": "Week6B_Waymo",
        "title": "Waymo: The Multi-Sensor Approach",
        "content_paragraphs": [
            "Waymo takes the opposite bet from Tesla: sensors everywhere. Their vehicles combine LIDAR, cameras, and radar for redundant perception. The philosophy is simple-when safety is paramount, redundancy provides a margin for error that single-modality systems cannot match.",

            "In October 2024, Waymo unveiled EMMA (End-to-End Multimodal Model for Autonomous Driving), powered by Google's Gemini large language model. EMMA directly maps raw camera sensor data into driving-specific outputs including planner trajectories, perception objects, and road graph elements. The system 'maximizes Gemini's world knowledge by representing non-sensor inputs and outputs as natural language text,' and has shown a 6.7% improvement in end-to-end planning performance.",

            "EMMA represents a convergence of computer vision and natural language processing. The model can provide 'an interpretable rationale for driving decisions'-explaining why it chose a particular path. But the approach has limitations: EMMA 'can process only a small amount of image frames, does not incorporate accurate 3D sensing modalities like LiDAR or radar, and is computationally expensive.'",

            "Waymo's safety record is the strongest argument for their approach. The company has driven over 127 million miles with only two fatal crashes, neither considered Waymo's fault. The autonomous system has driven 96 million 'rider-only' miles without any human driver. They claim their vehicles are 91% less likely to be involved in crashes resulting in serious injury compared to an average human driver over the same distance.",

            "As of late 2025, Waymo is delivering 450,000 paid rides per week with approximately 2,500 robotaxis in the United States-making it one of the largest driverless fleets in the world. The company is expanding into roughly 20 new markets, including first international launches in London and Tokyo, and is approaching a $100 billion valuation.",

            "The contrast with Tesla is stark. Waymo prioritizes safety and redundancy, operates in geofenced areas, and charges for rides. Tesla prioritizes scale and data collection, deploys everywhere, and sells vehicles. Both approaches could ultimately succeed-or fail. The experiment continues.",
        ],
        "key_facts": [
            ("Sensor approach", "LIDAR + cameras + radar"),
            ("EMMA model", "Built on Gemini LLM"),
            ("Total miles driven", "127+ million"),
            ("Rider-only miles", "96 million"),
            ("Safety improvement", "91% less likely to cause serious injury"),
            ("Weekly rides (2025)", "450,000"),
            ("Fleet size", "~2,500 robotaxis"),
            ("Valuation", "Approaching $100 billion"),
        ],
        "questions": [
            {
                "question": "Waymo uses LIDAR + cameras + radar while Tesla uses cameras only. Which company's approach would you invest in?",
                "notes": [
                    "Waymo: Redundancy and safety first, cost will decrease",
                    "Tesla: Scale data collection, software will improve",
                    "Waymo: Revenue-generating service now",
                    "Tesla: Consumer product, broader market"
                ]
            },
            {
                "question": "EMMA uses a language model (Gemini) for driving decisions. What are the implications of LLMs controlling vehicles?",
                "notes": [
                    "LLMs can reason about complex scenarios",
                    "But: Hallucination risk in safety-critical settings",
                    "Can provide interpretable rationale for decisions",
                    "Convergence of CV and NLP"
                ]
            },
            {
                "question": "Waymo has driven 127+ million miles with only 2 fatal crashes (neither their fault). How much data is 'enough' to prove safety?",
                "notes": [
                    "Human baseline: ~1 fatality per 100M miles",
                    "127M miles is good, but rare events are... rare",
                    "Statistical confidence requires billions of miles",
                    "Simulation helps but is not the same"
                ]
            },
            {
                "question": "Waymo operates robotaxis in limited areas. Tesla aims for everywhere. Which strategy makes more sense?",
                "notes": [
                    "Geofenced: Can map in detail, prove safety first",
                    "Everywhere: More data, broader applicability",
                    "Geofenced: Capital-intensive (dedicated fleet)",
                    "Everywhere: Leverages customer vehicles"
                ]
            }
        ],
        "key_insight": "Different companies make fundamentally different bets on the path to autonomy. Waymo prioritizes safety and redundancy; Tesla prioritizes scale and data.",
        "primary_source": "waymo.com/blog/2024/10/introducing-emma",
        "teaching_tips": [
            "Pair with Tesla case for great compare/contrast",
            "Good for discussing different AI strategies",
            "Connect to lecture: Sensor fusion, multi-task learning",
            "Ask: Which company's approach would you invest in?"
        ]
    },

    # WEEK 6C: JOHN DEERE
    {
        "week": 6,
        "topic": "Convolutional Neural Networks",
        "id": "Week6C_JohnDeere",
        "title": "John Deere See & Spray: Precision Agriculture",
        "content_paragraphs": [
            "In 2017, John Deere paid $305 million to acquire Blue River Technology, a startup founded by two Stanford graduate students in 2011 with a mission to 'make farming more sustainable through robotics and computer vision.' The acquisition brought See & Spray technology to the world's largest agricultural equipment company.",

            "The core innovation is elegantly simple: use computer vision to distinguish weeds from crops, then spray herbicide only on the weeds. Traditional spraying treats entire fields uniformly, wasting chemicals on crop plants that do not need them. See & Spray targets individual weeds with precision measured in fractions of a second.",

            "The latest version, See & Spray Ultimate, uses 36 cameras and processors mounted on a carbon-fiber boom to scan the ground at 2,500 square feet per second (traveling at 15 mph). The system 'identifies and hits the weed in less time than the blink of an eye.' Operating at up to 12 mph, it can reduce non-residual herbicide use by more than two-thirds while maintaining hit rates comparable to traditional spraying.",

            "John Deere estimates See & Spray reduces herbicide usage by 77%-some reports suggest up to 90% reduction is possible. For farmers, this means dramatically lower chemical costs. For the environment, it means less herbicide runoff into waterways. For John Deere, it means a compelling value proposition that justifies premium pricing on equipment.",

            "The product timeline shows deliberate development. By 2017, Blue River launched See & Spray for targeting individual weeds. Deere revealed the technology at CES 2020. In 2021, they introduced See & Spray Select for fallow fields. A year later came See & Spray Ultimate-a two-tank system for treating weeds in growing crops including corn, soybeans, and cotton.",

            "Beyond immediate sales, John Deere gains something arguably more valuable: data on millions of acres. Every pass of a See & Spray machine generates detailed imagery of crop and weed patterns. This data helps improve models faster than competitors without fleet-scale collection-creating a potential moat that extends beyond the equipment sale to an ongoing data relationship with farmers.",
        ],
        "key_facts": [
            ("Acquisition price", "$305 million (2017)"),
            ("Cameras on Ultimate", "36"),
            ("Scanning speed", "2,500 sq ft per second at 15 mph"),
            ("Herbicide reduction", "77% (up to 90% possible)"),
            ("Crops supported", "Corn, soybeans, cotton"),
            ("Data captured", "Detailed imagery of millions of acres"),
        ],
        "questions": [
            {
                "question": "John Deere paid $305M for Blue River. What were they buying-technology, talent, data, or market position?",
                "notes": [
                    "Technology: See & Spray system",
                    "Talent: Computer vision expertise",
                    "Data: Training data for agricultural CV",
                    "Position: Leadership in precision ag",
                    "All four, but data moat may be most valuable long-term"
                ]
            },
            {
                "question": "See & Spray reduces herbicide by 77%. Who captures the value-farmers, John Deere, or society?",
                "notes": [
                    "Farmers: Lower input costs",
                    "John Deere: Premium pricing on equipment",
                    "Society: Environmental benefits",
                    "Value distribution depends on market structure"
                ]
            },
            {
                "question": "The system must distinguish crops from weeds at 15 mph. What makes this computer vision problem challenging?",
                "notes": [
                    "Speed: Milliseconds to decide",
                    "Similarity: Weeds and crops look alike",
                    "Variability: Lighting, soil, growth stages",
                    "Consequences: Missing weeds or killing crops"
                ]
            },
            {
                "question": "John Deere now has detailed imagery of millions of acres. What are the strategic implications of this data?",
                "notes": [
                    "Improve models faster than competitors",
                    "Cross-sell other services (yield prediction)",
                    "Lock-in effect (leaving means losing data)",
                    "Potential for data marketplace"
                ]
            }
        ],
        "key_insight": "Computer vision transforms industries by making visual inspection scalable. The equipment sale is the beginning of a data relationship.",
        "primary_source": "agfundernews.com/breaking-exclusive-john-deere-acquires-see-spray-robotics-startup-blue-river-technology-305m",
        "teaching_tips": [
            "Good case for non-tech industry AI application",
            "Data moat discussion connects to strategy",
            "Connect to lecture: Real-time CNN inference",
            "Environmental angle can generate discussion"
        ]
    },

    # WEEK 7A: KLARNA
    {
        "week": 7,
        "topic": "Transformers & LLMs",
        "id": "Week7A_Klarna",
        "title": "Klarna's AI Reversal: When Automation Goes Too Far",
        "content_paragraphs": [
            "In February 2024, Klarna made headlines with a bold announcement: their AI assistant, developed in partnership with OpenAI, had handled 2.3 million customer conversations in its first month-work 'equivalent to 700 full-time agents.' The AI was operating in 35 languages, reducing resolution time dramatically. The fintech company seemed to have cracked AI customer service at scale.",

            "Between 2022 and 2024, Klarna cut approximately 700 customer service jobs and replaced them with AI-powered solutions. The company's workforce dropped from about 5,000 employees to 3,800. CEO Sebastian Siemiatkowski championed the transformation, positioning Klarna as a leader in enterprise AI adoption.",

            "Then the problems emerged. By early 2025, customer satisfaction had fallen sharply and service quality was inconsistent. Users reported experiences that were 'either impersonal or insufficient, particularly for nuanced issues.' The situation became so strained that Klarna was asking software engineers, designers, and marketing staff to help answer customer inquiries.",

            "CEO Siemiatkowski publicly acknowledged the failure: 'Cost unfortunately seems to have been a too predominant evaluation factor.' The AI-driven transition had 'negatively affected service and product quality.' Klarna's experience showed that 'implementing AI too hastily can compromise service quality, damage brand trust, and ultimately require costly reversals.'",

            "By spring 2025, Klarna reversed course. The company began rehiring human staff-remote support workers with flexible schedules, targeting students, parents, rural workers, and loyal Klarna users. The new model blends AI with human support: AI handles routine queries while humans remain available for complex issues and emotional support.",

            "The lesson extends beyond Klarna. As one analysis concluded: 'AI technology is yet to compete with the nuance and emotional capabilities of human workers, particularly in cases such as customer service, where a sense of empathy is needed.' Metrics that showed efficiency gains masked underlying quality problems that ultimately damaged the brand.",
        ],
        "key_facts": [
            ("AI launch", "February 2024 (partnership with OpenAI)"),
            ("First month conversations", "2.3 million"),
            ("Claimed equivalence", "'700 full-time agents'"),
            ("Jobs cut (2022-2024)", "~700 customer service positions"),
            ("Workforce reduction", "5,000 to 3,800 employees"),
            ("2025 reversal", "Rehiring human staff with hybrid model"),
        ],
        "questions": [
            {
                "question": "Klarna reported AI doing the work of 700 agents. What metrics might have hidden quality problems?",
                "notes": [
                    "Resolution time does not measure resolution quality",
                    "Conversation count does not measure satisfaction",
                    "Cost savings metric ignores customer experience",
                    "Short-term metrics can miss long-term damage"
                ]
            },
            {
                "question": "CEO Siemiatkowski admitted 'cost was too predominant a factor.' How should companies balance efficiency and quality in AI deployment?",
                "notes": [
                    "Track quality metrics alongside efficiency",
                    "Customer satisfaction, NPS, repeat contact rate",
                    "Pilot before full deployment",
                    "Human oversight for complex cases"
                ]
            },
            {
                "question": "Klarna now uses a hybrid model: AI for routine queries, humans for complex issues. How would you design the handoff?",
                "notes": [
                    "AI handles: Balance inquiries, payment dates, simple questions",
                    "Escalation triggers: Negative sentiment, repeated questions, complaints",
                    "Customer can always request human",
                    "Seamless transition preserving context"
                ]
            },
            {
                "question": "What does Klarna's reversal suggest about the current state of LLMs in customer service?",
                "notes": [
                    "LLMs lack empathy and nuance",
                    "Complex/emotional situations require humans",
                    "Technology is not ready for full replacement",
                    "Augmentation > replacement for now"
                ]
            }
        ],
        "key_insight": "The gap between 'demo' and 'production' is real. Customer-facing AI requires careful monitoring of quality, not just efficiency metrics.",
        "primary_source": "tech.co/news/klarna-reverses-ai-overhaul",
        "teaching_tips": [
            "Very current case (2024-2025)-students may have seen headlines",
            "Good for discussing metric selection and Goodhart's Law",
            "Connect to lecture: LLM capabilities and limitations",
            "Ask: Have you had a frustrating chatbot experience?"
        ]
    },

    # WEEK 7B: JPMORGAN COIN (NLP version for Week 7)
    {
        "week": 7,
        "topic": "Transformers & LLMs",
        "id": "Week7B_JPMorgan_COIN",
        "title": "JPMorgan COIN: NLP for Contract Intelligence",
        "content_paragraphs": [
            "Before COIN, JPMorgan's lawyers and loan officers faced a crushing burden: 360,000 hours annually spent reviewing commercial loan agreements. That is more than 41 years of continuous work devoted to what one executive called a 'mind-numbing job' of interpreting dense legal language, identifying key terms, and flagging potential risks.",

            "COIN (Contract Intelligence) transformed this process using natural language processing. The system reviews documents in seconds, deriving '150 attributes from 12,000 commercial credit agreements and contracts' almost instantly. Unlike human reviewers, 'it is less error-prone and never asks for vacation.'",

            "The technical approach uses unsupervised learning to identify and categorize repeated clauses in credit contracts. This allows COIN to discover patterns in legal language without requiring massive labeled datasets-a crucial advantage when dealing with evolving contract types and emerging clause structures.",

            "Quality improved alongside speed. COIN 'has helped JPMorgan cut down on loan-servicing mistakes, most of which stemmed from human error in interpreting 12,000 new wholesale contracts per year.' The algorithm proved more accurate than human lawyers at consistent interpretation, making the investment about quality as much as cost savings.",

            "JPMorgan planned to expand COIN to more complex filings-credit-default swaps and custody agreements-and eventually to interpret entirely new regulations. The vision: AI that can read and understand legal language at scale, freeing human experts for judgment calls and relationship management.",

            "The liability question remains unresolved. When AI misses a critical contract term, who bears responsibility? Currently, lawyers still sign off on AI-reviewed documents. But as systems like COIN handle increasingly complex analysis, the legal framework for AI-assisted professional work continues to evolve.",
        ],
        "key_facts": [
            ("Previous manual effort", "360,000 hours annually"),
            ("Processing time with COIN", "Seconds"),
            ("Early capability", "150 attributes from 12,000 contracts"),
            ("Learning approach", "Unsupervised (discovers clause patterns)"),
            ("Quality impact", "Reduced loan-servicing mistakes"),
            ("Future expansion", "Credit-default swaps, custody agreements, new regulations"),
        ],
        "questions": [
            {
                "question": "COIN uses NLP to interpret legal language. What makes legal documents particularly challenging for language models?",
                "notes": [
                    "Precise language where small differences matter",
                    "Domain-specific terminology",
                    "Nested references and cross-references",
                    "Consequences of misinterpretation are severe"
                ]
            },
            {
                "question": "COIN uses unsupervised learning to categorize clauses. Why unsupervised rather than supervised?",
                "notes": [
                    "No pre-existing labels for all clause types",
                    "Contracts evolve-new clause types emerge",
                    "Discovers patterns humans might miss",
                    "Supervised would require massive labeling effort"
                ]
            },
            {
                "question": "COIN reduced human errors in contract interpretation. What types of errors are humans prone to that AI might avoid?",
                "notes": [
                    "Fatigue-related mistakes (hour 360,000 vs. hour 1)",
                    "Inconsistent interpretation across reviewers",
                    "Missing patterns across large document sets",
                    "But AI might miss context humans catch"
                ]
            },
            {
                "question": "JPMorgan plans to use COIN for new regulations. How might NLP help interpret emerging legal requirements?",
                "notes": [
                    "Extract key obligations from regulatory text",
                    "Compare to existing policies/practices",
                    "Flag conflicts or gaps",
                    "Speed regulatory response"
                ]
            }
        ],
        "key_insight": "NLP can automate document-heavy professional work, but legal consequences require human oversight. AI augments rather than replaces expert judgment.",
        "primary_source": "abajournal.com/news/article/jpmorgan_chase_uses_tech_to_save_360000_hours_of_annual_work_by_lawyers_and",
        "teaching_tips": [
            "Good for discussing NLP applications in professional services",
            "Connect to lecture: How transformers process document structure",
            "Ask: Would you trust AI to review a contract you are signing?",
            "Can discuss automation of knowledge work broadly"
        ]
    },

    # WEEK 8A: MORGAN STANLEY
    {
        "week": 8,
        "topic": "Generative AI & Agents",
        "id": "Week8A_MorganStanley",
        "title": "Morgan Stanley: Enterprise AI at Scale",
        "content_paragraphs": [
            "In March 2023, Morgan Stanley became the first major Wall Street firm to deploy a bespoke solution based on GPT-4 in employees' hands. Called the AI @ Morgan Stanley Assistant, the tool gives financial advisors instant access to what the firm calls its 'intellectual capital'-a database of approximately 100,000 research reports and documents.",

            '"We have a knowledge base of over 100,000 documents and the idea is to bring that knowledge base-that differentiated intellectual capital-and put it together on top of the models of GPT-4," explained Jeff McMillan, head of analytics, data and innovation at Morgan Stanley wealth management. The goal: create "an ecosystem that is private to Morgan Stanley" making expert content available "24 hours a day 7 days a week virtually instantaneously."',

            "Importantly, Morgan Stanley's solutions do not use ChatGPT, which leverages GPT-3.5 and generates responses from the public internet. Instead, they use GPT-4 to generate responses exclusively from internal Morgan Stanley content, with appropriate controls. The bank spent months curating documents and using human experts to test responses before rolling out to 300 advisors in initial testing.",

            "The adoption has been remarkable. Today, over 98% of advisor teams actively use the AI Assistant. Access to documents jumped from 20% to 80%, dramatically reducing search time and increasing document retrieval efficiency. The barrier between knowledge and communication, as one managing director put it, has essentially disappeared.",

            "Building on this success, Morgan Stanley launched additional AI tools. AI @ Morgan Stanley Debrief uses Whisper and GPT-4 to turn Zoom recordings (with client consent) into actionable outputs-client notes automatically integrated into CRM systems and draft follow-ups summarizing key action items. AskResearchGPT serves Investment Banking, Sales & Trading, and Research with access to 70,000+ proprietary reports published annually.",

            "But questions remain about knowledge currency. Markets change faster than documents update. Different analysts may have conflicting views. The AI might present outdated analysis as current insight. And if junior analysts previously synthesized research-learning the business in the process-what is their development path now?",
        ],
        "key_facts": [
            ("Launch", "March 2023 (first major Wall Street firm with GPT-4)"),
            ("Model", "GPT-4 (not ChatGPT/GPT-3.5)"),
            ("Knowledge base", "100,000+ internal documents"),
            ("Advisor adoption", "98% of teams actively using"),
            ("Document access improvement", "20% to 80%"),
            ("Additional tools", "Debrief (meeting summaries), AskResearchGPT"),
            ("Research volume", "70,000+ proprietary reports annually"),
        ],
        "questions": [
            {
                "question": "Morgan Stanley chose GPT-4 with internal documents only-not ChatGPT with internet access. Why does this distinction matter?",
                "notes": [
                    "Regulatory requirements: Data cannot leave firm",
                    "Accuracy: Internal documents are vetted",
                    "Liability: Can trace sources",
                    "Competitive advantage: Proprietary insights"
                ]
            },
            {
                "question": "Document access jumped from 20% to 80% with the AI Assistant. What does this suggest about knowledge utilization before AI?",
                "notes": [
                    "Vast knowledge was essentially inaccessible",
                    "Search was too time-consuming",
                    "Advisors did not know what existed",
                    "AI unlocks dormant organizational knowledge"
                ]
            },
            {
                "question": "Markets change faster than documents update. How should the system handle outdated or conflicting information?",
                "notes": [
                    "Timestamp awareness and recency weighting",
                    "Surface conflicting analyst views",
                    "Human verification for time-sensitive decisions",
                    "Clear provenance for all answers"
                ]
            },
            {
                "question": "Junior analysts previously synthesized research as part of their training. How does AI change their development path?",
                "notes": [
                    "Less 'summarizing' work available",
                    "More 'validating AI output' work",
                    "Need new ways to learn the business",
                    "Shift to judgment and relationship skills earlier"
                ]
            }
        ],
        "key_insight": "Enterprise AI requires data strategy, integration, and governance-technology is the easy part. Competitive advantage comes from proprietary data and integration depth.",
        "primary_source": "cnbc.com/2023/09/18/morgan-stanley-chatgpt-financial-advisors.html",
        "teaching_tips": [
            "Good capstone case-brings together many course themes",
            "Discuss enterprise AI challenges vs. consumer AI",
            "Connect to lecture: RAG, fine-tuning, enterprise deployment",
            "Ask: What would you want an AI assistant to help with?"
        ]
    },

    # WEEK 8B: GITHUB COPILOT
    {
        "week": 8,
        "topic": "Generative AI & Agents",
        "id": "Week8B_Copilot",
        "title": "GitHub Copilot: AI Pair Programming",
        "content_paragraphs": [
            "GitHub collaborated with researchers to conduct what may be the most rigorous study of AI coding assistance: a randomized controlled trial (RCT) measuring Copilot's impact on developer productivity. The results were striking.",

            "Developers with access to GitHub Copilot completed tasks 55.8% faster than those without. The treatment group took on average 1 hour and 11 minutes to complete a coding task, while the control group took 2 hours and 41 minutes. These results were statistically significant (P=.0017), with a 95% confidence interval for the speed gain ranging from 21% to 89%.",

            "In a larger enterprise study with Accenture, GitHub found additional benefits beyond raw speed. Developers reported 90% higher job fulfillment when using Copilot, and 95% said they enjoyed coding more with Copilot's help. The study observed an 84% increase in successful builds, suggesting Copilot not only accelerates coding but may improve initial code quality.",

            "A notable finding: less experienced developers benefited most from Copilot. This has profound implications for training and hiring. If AI can help junior developers produce senior-level output, the value of memorizing syntax decreases while the value of problem decomposition and architectural thinking increases.",

            "The intellectual property questions are thorny. Copilot was trained on public GitHub repositories, raising concerns about whether AI-generated code might infringe on training data licenses. Active lawsuits are pending. Companies are developing policies for when and how AI-generated code can be used.",

            "Not all research is positive. A contrasting study from Uplevel Data Labs found that 'developers with Copilot access saw a significantly higher bug rate while their issue throughput remained consistent.' This suggests Copilot may have trade-offs: faster initial coding but potentially more debugging later. The full picture of AI-assisted development is still emerging.",
        ],
        "key_facts": [
            ("RCT result", "55.8% faster task completion"),
            ("Time comparison", "1hr 11min (Copilot) vs. 2hr 41min (control)"),
            ("Statistical significance", "P=.0017"),
            ("Job satisfaction", "90% higher fulfillment"),
            ("Coding enjoyment", "95% enjoyed coding more"),
            ("Build success", "84% increase in successful builds"),
            ("Who benefits most", "Less experienced developers"),
        ],
        "questions": [
            {
                "question": "Less experienced developers benefited most from Copilot. What does this imply for training and hiring?",
                "notes": [
                    "AI as learning accelerator",
                    "Less emphasis on syntax memorization",
                    "More emphasis on problem decomposition",
                    "Hiring: Look for reasoning, not rote coding"
                ]
            },
            {
                "question": "Copilot was trained on public code. What are the IP implications of AI-generated code?",
                "notes": [
                    "Does output infringe on training data licenses?",
                    "Who owns AI-generated code?",
                    "Active lawsuits pending",
                    "Companies developing policies"
                ]
            },
            {
                "question": "One study found higher bug rates with Copilot. How do you reconcile this with the productivity gains?",
                "notes": [
                    "Speed vs. quality trade-off",
                    "Faster initial coding, more debugging later",
                    "Net effect may depend on use case",
                    "Highlights need for testing/review"
                ]
            },
            {
                "question": "Developers report 95% higher enjoyment with Copilot. What aspects of coding become more/less enjoyable with AI assistance?",
                "notes": [
                    "Less enjoyable: Gone is tedious boilerplate",
                    "More enjoyable: Focus on interesting problems",
                    "More enjoyable: Reduced context switching",
                    "Risk: May reduce deep learning of fundamentals"
                ]
            }
        ],
        "key_insight": "AI tools change the skill mix required-less syntax memorization, more architecture and problem-solving. Rigorous evaluation (RCT) provides credible evidence of impact.",
        "primary_source": "github.blog/news-insights/research/research-quantifying-github-copilots-impact-on-developer-productivity-and-happiness/",
        "teaching_tips": [
            "Many CS students will have used Copilot",
            "Good for discussing AI's impact on knowledge work",
            "Connect to lecture: Code as language, LLMs for code",
            "RCT methodology is notable-discuss what makes evidence credible"
        ]
    },

    # WEEK 8C: NIKE
    {
        "week": 8,
        "topic": "Generative AI & Agents",
        "id": "Week8C_Nike",
        "title": "Nike: AI-Driven Digital Transformation",
        "content_paragraphs": [
            "Between 2018 and 2021, Nike acquired four AI companies in rapid succession: Zodiac (customer lifetime value modeling, March 2018), Celect (predictive inventory optimization, August 2019), Invertex (3D foot scanning), and Datalogue (data integration, February 2021). The acquisitions signaled Nike's aggressive push into data-driven retail.",

            "Each acquisition served a specific purpose. Zodiac provided insights on 'the value of an individual customer to enhance revenue and retention with the correct marketing, recommendations, and offers.' Celect brought hyperlocal demand sensing for inventory optimization. Datalogue's machine learning technology automated data preparation and integration-the plumbing needed to make the other systems work together.",

            "Early results looked promising. Invertex became the technology behind Nike Fit, the sizing app experience. Management credited Celect's technology for driving 100% year-over-year digital revenue growth in North America while lowering digital fulfillment costs per unit. Nike Direct grew from 10% to 26% of total revenue.",

            "Then came 2024. Nike confronted what analysts called a 'big failure due to poor digital transformation initiatives.' The company's market value dropped $25 billion in a single day. What went wrong?",

            "Nike had 'over-rotated' on digital and data at the expense of domain expertise. The company adopted a one-size-fits-all approach, eliminating 70% of category experts and replacing them with generalized AI/ML professionals. As one analysis noted: 'The over-dependence on data analytics and the complete removal of domain expertise led to Nike's digital transformation failing to meet real-world customer expectations.'",

            "The focus on digital also damaged wholesale relationships. Nike pulled back from retailers like Foot Locker to push direct-to-consumer sales. This increased overdiscounting and compressed profit margins from 46% to 43.5%. Digital sales declined in fiscal year 2024, forcing Nike to rebuild the wholesale partnerships it had abandoned. The AI transformation delivered metrics but created strategic problems that metrics could not capture.",
        ],
        "key_facts": [
            ("AI acquisitions", "4 companies (2018-2021)"),
            ("Zodiac", "March 2018 (customer lifetime value)"),
            ("Celect", "August 2019 (inventory optimization)"),
            ("Datalogue", "February 2021 (data integration)"),
            ("Nike Direct growth", "10% to 26% of revenue"),
            ("2024 market cap loss", "$25 billion in one day"),
            ("Category experts cut", "70%"),
            ("Margin compression", "46% to 43.5%"),
        ],
        "questions": [
            {
                "question": "Nike acquired 4 AI companies in 3 years. When does it make sense to buy vs. build AI capabilities?",
                "notes": [
                    "Buy: Speed to market, acquire talent, proven technology",
                    "Build: Unique needs, core competency, control",
                    "Nike bought because retail AI was not core competency",
                    "Integration is the hard part"
                ]
            },
            {
                "question": "Nike eliminated 70% of category experts while hiring AI/ML professionals. What went wrong with this trade-off?",
                "notes": [
                    "Data cannot capture everything (taste, trends, relationships)",
                    "Domain expertise provides judgment AI lacks",
                    "Over-indexed on what is measurable",
                    "Lost institutional knowledge"
                ]
            },
            {
                "question": "Nike Direct grew from 10% to 26% of revenue, but the strategy ultimately failed. What metrics might have hidden the problems?",
                "notes": [
                    "Digital revenue growth looked good",
                    "But: Damaged wholesale relationships",
                    "But: Increased discounting",
                    "But: Lost retail presence and brand visibility"
                ]
            },
            {
                "question": "Nike is now rebuilding wholesale partnerships it abandoned. What does this reversal teach about AI transformation?",
                "notes": [
                    "Technology changes faster than relationships rebuild",
                    "Burning bridges has long-term costs",
                    "Balance innovation with existing strengths",
                    "Not everything should be optimized"
                ]
            }
        ],
        "key_insight": "AI strategy is not just about technology-it is about integrating capabilities without losing domain expertise. Success metrics can mask strategic problems.",
        "primary_source": "techtidesolutions.com/blog/nike-digital-transformation/",
        "teaching_tips": [
            "Good closing case-shows both success and failure",
            "Connects AI to broader business strategy",
            "Nike brand is familiar to all students",
            "Ask: What would you have done differently?"
        ]
    },
]

# ============================================
# GENERATE ALL HANDOUTS
# ============================================

import os

# Create directories if needed
os.makedirs("/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/student", exist_ok=True)
os.makedirs("/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/instructor", exist_ok=True)

for case in cases:
    # Student handout (Word document)
    student_file = f"/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/student/{case['id']}.docx"
    create_student_handout(
        student_file,
        case['week'],
        case['topic'],
        case['title'],
        case['content_paragraphs'],
        case['key_facts'],
        case['questions'],
        case['primary_source']
    )

    # Instructor notes (Word document)
    instructor_file = f"/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/instructor/{case['id']}_INSTRUCTOR.docx"
    create_instructor_notes(
        instructor_file,
        case['week'],
        case['topic'],
        case['title'],
        case['questions'],
        case['key_insight'],
        case.get('teaching_tips', [])
    )

print("\n All handouts generated successfully!")
print(" Student handouts: /sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/student/")
print(" Instructor notes: /sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/instructor/")
