"""
Generate Three New Aligned Case Studies for MGMT298D
UCLA Anderson - Science and Strategy of AI

New cases that align with notebook applications:
- Week 1B: H&M Demand Forecasting (aligns with H&M sales notebook)
- Week 2B: Carvana/CarMax Vehicle Pricing (aligns with Range Rover pricing notebook)
- Week 5B: Stitch Fix AI Styling (aligns with Fashion MNIST notebook)

All facts are drawn from verified sources.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# UCLA Colors (RGB)
UCLA_BLUE = RGBColor(0x27, 0x74, 0xAE)
UCLA_DARK_BLUE = RGBColor(0x00, 0x55, 0x87)
UCLA_GOLD = RGBColor(0xFF, 0xD1, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)


def add_horizontal_line(paragraph, color_hex='FFD100'):
    """Add a horizontal line after a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_student_handout(filename, week_num, week_topic, case_title, content_paragraphs, key_facts, questions, primary_source):
    """Create student-facing handout as Word document"""
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("MGMT298D: Science and Strategy of AI | UCLA Anderson")
    run.font.size = Pt(10)
    run.font.color.rgb = UCLA_BLUE
    run.font.name = 'Arial'

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(case_title)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Week {week_num}: {week_topic}")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Arial'

    add_horizontal_line(subtitle, 'FFD100')

    section_header = doc.add_paragraph()
    run = section_header.add_run("The Situation")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    section_header.space_before = Pt(6)
    section_header.space_after = Pt(3)

    for para in content_paragraphs:
        p = doc.add_paragraph()
        if para.startswith('"') or para.startswith('"'):
            run = p.add_run(para)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
        else:
            run = p.add_run(para)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    section_header = doc.add_paragraph()
    run = section_header.add_run("Discussion Questions")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    section_header.space_before = Pt(6)
    section_header.space_after = Pt(3)

    for i, q in enumerate(questions, 1):
        q_num = doc.add_paragraph()
        run = q_num.add_run(f"Question {i}")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_BLUE
        run.font.name = 'Arial'
        q_num.space_before = Pt(4)
        q_num.space_after = Pt(1)

        q_text = doc.add_paragraph()
        run = q_text.add_run(q['question'])
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Arial'
        q_text.paragraph_format.left_indent = Inches(0.25)
        q_text.space_after = Pt(3)

    footer_line = doc.add_paragraph()
    footer_line.space_before = Pt(6)
    add_horizontal_line(footer_line, 'CCCCCC')

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Source: {primary_source}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Arial'

    doc.save(filename)
    print(f"Created: {filename}")


def create_instructor_notes(filename, week_num, week_topic, case_title, questions, key_insight, teaching_tips):
    """Create instructor-only teaching notes as Word document"""
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("INSTRUCTOR NOTES - DO NOT DISTRIBUTE")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run.font.name = 'Arial'

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(case_title)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Week {week_num}: {week_topic}")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Arial'

    add_horizontal_line(subtitle, 'C62828')

    insight_header = doc.add_paragraph()
    run = insight_header.add_run("Key Insight (Goal of Discussion)")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run.font.name = 'Arial'
    insight_header.space_before = Pt(12)
    insight_header.space_after = Pt(6)

    insight_box = doc.add_paragraph()
    insight_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = insight_box.add_run(key_insight)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    insight_box.space_after = Pt(10)

    guide_header = doc.add_paragraph()
    run = guide_header.add_run("Discussion Guide")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = UCLA_DARK_BLUE
    run.font.name = 'Arial'
    guide_header.space_before = Pt(10)
    guide_header.space_after = Pt(6)

    for i, q in enumerate(questions, 1):
        q_para = doc.add_paragraph()
        run = q_para.add_run(f"Question {i}: {q['question']}")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = UCLA_BLUE
        run.font.name = 'Arial'
        q_para.space_before = Pt(8)
        q_para.space_after = Pt(4)

        notes_label = doc.add_paragraph()
        run = notes_label.add_run("Teaching Notes:")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Arial'
        notes_label.paragraph_format.left_indent = Inches(0.25)
        notes_label.space_after = Pt(2)

        for note in q['notes']:
            note_para = doc.add_paragraph()
            run = note_para.add_run(f"  {note}")
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
            note_para.paragraph_format.left_indent = Inches(0.25)
            note_para.space_after = Pt(2)

    if teaching_tips:
        tips_header = doc.add_paragraph()
        run = tips_header.add_run("Additional Teaching Tips")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        run.font.name = 'Arial'
        tips_header.space_before = Pt(12)
        tips_header.space_after = Pt(6)

        for tip in teaching_tips:
            tip_para = doc.add_paragraph()
            run = tip_para.add_run(f"  {tip}")
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Arial'
            tip_para.paragraph_format.left_indent = Inches(0.25)
            tip_para.space_after = Pt(4)

    doc.save(filename)
    print(f"Created: {filename}")


# ============================================
# NEW CASE DATA - Aligned with Notebooks
# ============================================

new_cases = [
    # ============================================
    # WEEK 1B: H&M DEMAND FORECASTING
    # Aligns with Week 1 notebook (H&M sales data)
    # ============================================
    {
        "week": 1,
        "topic": "Linear Regression",
        "id": "Week1B_HM",
        "title": "H&M: AI-Powered Demand Forecasting",
        "content_paragraphs": [
            "H&M, the world's second-largest fashion retailer, faced a crisis familiar to the industry: billions of dollars in unsold inventory. In 2018, the company reported $4.3 billion worth of unsold clothes. The fast fashion model that had driven growth was now creating massive waste-and massive losses from markdowns.",

            "The company responded by building one of fashion's most ambitious AI initiatives. H&M created a dedicated AI department in 2018 with a clear mandate: use machine learning to predict what customers will want, where they will want it, and when. Today, that team has grown to 270 data scientists working to apply AI across the company.",

            "The technical approach centers on time-series forecasting enhanced with neural networks. H&M's models are trained on more than 10 years of historical sales data, combined with external signals including weather patterns, local events, and social media trends. The system forecasts item-level demand per store-down to specific sizes and colors-and retrains daily to capture emerging patterns.",

            "The partnership with Google Cloud proved essential for scaling. H&M built a centralized data platform that integrated information from sales channels, warehouses, and suppliers, providing what the company calls 'a holistic view of the entire supply chain.' This infrastructure enables the ML models to run at the scale required for a retailer with 4,000+ stores worldwide.",

            "The results have been substantial. Forecasting accuracy reached 85-90%, and overstock dropped by 25%-freeing up millions in working capital. Stockouts during high-demand periods like Black Friday fell by 14%. Sell-through rates climbed 15-20%, reducing the need for markdowns that had been eroding margins.",

            "But the system required more than algorithms. H&M restructured how stores receive inventory, moving from standardized assortments to locally-tailored selections based on regional demand signals. The AI does not just predict-it reshapes operations. As one H&M executive noted: 'Without precise demand planning, the result is waste-a historical problem for the fashion industry.'",
        ],
        "key_facts": [
            ("Unsold inventory (2018)", "$4.3 billion"),
            ("AI team size", "270 data scientists"),
            ("Training data", "10+ years of historical sales"),
            ("Forecast granularity", "Item-level per store (size, color)"),
            ("Forecast accuracy", "85-90%"),
            ("Overstock reduction", "25%"),
            ("Stockout reduction", "14% during peak periods"),
        ],
        "questions": [
            {
                "question": "H&M trains on 10+ years of historical data but retrains daily. Why is this combination of long history and frequent updates important for demand forecasting?",
                "notes": [
                    "Long history captures seasonal patterns (winter coats, summer dresses)",
                    "Daily retraining catches emerging trends and anomalies",
                    "Fashion moves fast-last year's hot item may be dead",
                    "External shocks (weather, events) need rapid incorporation"
                ]
            },
            {
                "question": "H&M forecasts demand at the item-store-size-color level. What are the trade-offs between granular and aggregate forecasting?",
                "notes": [
                    "Granular: More actionable but harder (sparse data per cell)",
                    "Aggregate: More stable but less useful operationally",
                    "The 'long tail' problem: Many SKUs sell few units",
                    "Hierarchy helps: Forecast at multiple levels, reconcile"
                ]
            },
            {
                "question": "Overstock fell 25% but H&M still has unsold inventory. Why can demand forecasting never be perfect in fashion?",
                "notes": [
                    "Fashion is inherently unpredictable (trends, viral moments)",
                    "Long lead times: Must order months before selling",
                    "Substitution effects: If item A sells out, does B sell more?",
                    "Forecast accuracy of 85-90% means 10-15% error remains"
                ]
            },
            {
                "question": "H&M moved from standardized store assortments to locally-tailored ones. What organizational changes does AI-driven forecasting require?",
                "notes": [
                    "Store managers lose some autonomy (algorithm decides)",
                    "Supply chain must handle more complexity (different mix per store)",
                    "Buyers must trust the model over intuition",
                    "Need new KPIs: forecast accuracy, not just sales"
                ]
            }
        ],
        "key_insight": "Demand forecasting is not just about building better models-it requires reshaping operations to act on predictions. The algorithm is only valuable if the organization can respond.",
        "primary_source": "https://www.supplychaindive.com/news/HM-AI-supply-chain-sustainable/570400/",
        "teaching_tips": [
            "Direct connection to Week 1 notebook-students will work with H&M data",
            "Ask: What features would you include beyond price and lag sales?",
            "Good discussion on forecast granularity vs. data sparsity",
            "Can discuss sustainability angle-less waste from better forecasting"
        ]
    },

    # ============================================
    # WEEK 2B: CARVANA/CARMAX VEHICLE PRICING
    # Aligns with Week 2 notebook (Range Rover pricing)
    # ============================================
    {
        "week": 2,
        "topic": "Tree-Based Predictions",
        "id": "Week2B_Carvana",
        "title": "Carvana & CarMax: Algorithmic Vehicle Pricing",
        "content_paragraphs": [
            "When you sell a car to Carvana or CarMax, an algorithm decides what to pay you-often within seconds. These companies have built sophisticated machine learning systems that analyze millions of vehicle transactions to generate instant valuations. The models must balance competing objectives: pay enough to acquire inventory, but not so much that resale becomes unprofitable.",

            "CarMax has been using AI and machine learning for vehicle pricing for over 12 years. As their technology leadership explains: 'Everything CarMax does is driven by algorithms that run constantly, looking at the market and setting prices for what they should be paying for trade-ins as well as what their inventory should be priced at.' The company processes real-time market data, competitor prices, and consumer behavior to recommend optimal pricing strategies.",

            "Carvana's approach centers on their proprietary pricing algorithm, which powers tools like the Carvana Value Tracker. The system uses gradient boosting machines and neural networks to handle complex, nonlinear relationships in vehicle pricing. Models are trained on millions of wholesale transactions, considering factors like sale price, odometer reading, vehicle condition, and regional demand variations.",

            "The accuracy claims are impressive. Carvana reports their Value Tracker's predictions typically fall within 2% of actual transaction prices. The system continuously learns from new data, becoming more accurate over time. Regional demand patterns matter significantly-high-demand urban areas can see vehicle values elevated by up to 5% compared to regions with declining demand.",

            "But algorithmic pricing creates challenges. During the pandemic-era used car boom, these algorithms struggled with unprecedented market conditions. Carvana's aggressive purchasing-driven by models trained on different market dynamics-contributed to significant losses when prices normalized. The same algorithms that enabled scale also amplified mistakes.",

            "CarMax has invested in computer vision through a partnership with UVeye, using AI to automate vehicle condition assessment. Cameras scan the body, tires, and undercarriage of vehicles, generating standardized condition reports. This addresses a key input to pricing models: reliable, consistent data about vehicle quality that does not depend on human inspection variability.",
        ],
        "key_facts": [
            ("CarMax ML history", "12+ years of AI/ML for pricing"),
            ("Carvana accuracy claim", "Within 2% of transaction prices"),
            ("Model types", "Gradient boosting, neural networks"),
            ("Key features", "Price, odometer, condition, regional demand"),
            ("Regional price variation", "Up to 5% based on local demand"),
            ("CarMax partnership", "UVeye for AI vehicle inspection"),
        ],
        "questions": [
            {
                "question": "CarMax and Carvana both use tree-based models for pricing. Why might gradient boosting outperform linear regression for used car valuation?",
                "notes": [
                    "Non-linear relationships (mileage effect is not constant)",
                    "Interactions matter (luxury + high mileage vs. economy + high mileage)",
                    "Categorical variables (make, model, trim) handled naturally",
                    "Can capture threshold effects (warranty cutoffs, etc.)"
                ]
            },
            {
                "question": "Carvana's algorithms struggled during the pandemic car boom. What does this reveal about the limits of ML pricing models?",
                "notes": [
                    "Models assume the future resembles the past",
                    "Unprecedented demand broke historical patterns",
                    "No training data for 30%+ price spikes",
                    "Human judgment needed for regime changes"
                ]
            },
            {
                "question": "CarMax uses computer vision to assess vehicle condition. How does this improve pricing model inputs compared to human inspection?",
                "notes": [
                    "Consistency: Same standards across all locations",
                    "Speed: Seconds vs. minutes for inspection",
                    "Documentation: Photos provide audit trail",
                    "But: May miss things humans catch (smell, sounds)"
                ]
            },
            {
                "question": "If you were building a car pricing model, what features would you include? Which would be hardest to obtain?",
                "notes": [
                    "Easy: Year, make, model, mileage, location",
                    "Medium: Condition score, service history, accident reports",
                    "Hard: Current owner's urgency to sell, local competitor inventory",
                    "Impossible: What a specific buyer will pay (idiosyncratic value)"
                ]
            }
        ],
        "key_insight": "Tree-based models excel at pricing because they capture non-linear relationships and interactions. But models trained on historical data fail when market regimes change.",
        "primary_source": "https://www.modernretail.co/technology/carmaxs-top-tech-exec-shares-his-keys-to-reinventing-a-legacy-retailer-in-the-age-of-ai/",
        "teaching_tips": [
            "Direct connection to Week 2 notebook-students price Range Rovers",
            "Ask: Which features in the notebook are most important? Compare to case.",
            "Good discussion on model interpretability (can explain why price is X)",
            "Carvana's pandemic losses echo Zillow-market regime changes break models"
        ]
    },

    # ============================================
    # WEEK 5B: STITCH FIX AI STYLING
    # Aligns with Week 5A notebook (Fashion MNIST)
    # ============================================
    {
        "week": 5,
        "topic": "Neural Networks",
        "id": "Week5B_StitchFix",
        "title": "Stitch Fix: When Algorithms Meet Stylists",
        "content_paragraphs": [
            "Stitch Fix was built on a bet that neither humans nor algorithms alone could solve personal styling-but together, they could. The company pairs machine learning with human stylists to send customers personalized clothing selections. As founder Katrina Lake put it: 'A good person plus a good algorithm is far superior to the best person or the best algorithm alone.'",

            "The data foundation is massive. Stitch Fix collects 90 data points from each new customer through an onboarding survey covering style preferences, size, fit, budget, and lifestyle. Beyond surveys, the company has accumulated over 4.5 billion textual data points from customer feedback-'more than all of Wikipedia,' according to their data science team. Every thumbs-up, written note, and Pinterest board contributes to the learning.",

            "The core algorithm predicts 'probability of purchase' by scoring each item against each customer. The model ingests new data daily from Style Shuffle (where customers swipe on items), written feedback, and purchase history. But Stitch Fix goes beyond collaborative filtering-they analyze photos of clothing using computer vision to understand style attributes that are hard to describe in words.",

            "The scale is remarkable. Stitch Fix generates 13 million new outfit combinations every day and shows approximately 43 million outfit suggestions to clients daily. Their Outfit Creation Model, trained on millions of outfits assembled by human stylists, considers real-time inventory and individual preferences to create coherent looks-not just individual items.",

            "Human stylists remain central. While algorithms generate initial recommendations, stylists have the power to override any suggestion. They handle nuanced requests that algorithms struggle with-'I need a dress for an outdoor wedding in July' requires understanding context that pure pattern-matching misses. Stitch Fix even matches stylists to clients based on 'latent style' similarity.",

            "The company uses GPT-4 to help stylists process customer history. For loyal clients who have shared hundreds of notes over years, AI generates concise summaries of preferences, dislikes, and fit information. This augments rather than replaces human judgment-stylists make final decisions armed with AI-processed context.",
        ],
        "key_facts": [
            ("Onboarding data points", "90 per customer"),
            ("Total textual data", "4.5 billion data points"),
            ("Daily outfit combinations generated", "13 million"),
            ("Daily outfit suggestions shown", "43 million"),
            ("Model training data", "Millions of stylist-created outfits"),
            ("AI assistance", "GPT-4 for summarizing client history"),
        ],
        "questions": [
            {
                "question": "Stitch Fix uses both collaborative filtering and computer vision. Why might visual analysis of clothing be necessary beyond just tracking what people buy?",
                "notes": [
                    "Purchase data is sparse (few items per person)",
                    "Style has visual dimensions hard to describe in text",
                    "Can find 'similar' items that were never purchased together",
                    "Enables recommendations for new items with no purchase history"
                ]
            },
            {
                "question": "Stitch Fix lets stylists override algorithm recommendations. When should humans override ML systems, and when should they trust them?",
                "notes": [
                    "Override for: Context-dependent requests, edge cases, emotional factors",
                    "Trust for: Large-scale patterns, consistent preferences, data-rich decisions",
                    "Key insight: Algorithms optimize for average, humans handle exceptions",
                    "Need feedback loop: Track when overrides improve outcomes"
                ]
            },
            {
                "question": "Stitch Fix matches stylists to clients based on 'latent style' similarity. What is a latent representation, and why is it useful?",
                "notes": [
                    "Latent = hidden representation learned by neural network",
                    "Encodes style in numbers (vectors) that capture similarity",
                    "Two clients with similar latent style like similar things",
                    "Enables matching without explicit rules"
                ]
            },
            {
                "question": "Stitch Fix uses GPT-4 to summarize years of customer notes for stylists. What are the risks and benefits of AI-generated summaries?",
                "notes": [
                    "Benefit: Saves stylist time, surfaces forgotten preferences",
                    "Benefit: Consistent processing of unstructured text",
                    "Risk: May miss nuance or context",
                    "Risk: Hallucination could introduce false information",
                    "Mitigation: Stylist reviews and can access original notes"
                ]
            }
        ],
        "key_insight": "Neural networks can learn latent representations of style that capture patterns humans cannot articulate. But combining AI with human judgment outperforms either alone.",
        "primary_source": "https://algorithms-tour.stitchfix.com/",
        "teaching_tips": [
            "Connection to Week 5A notebook-students classify fashion images",
            "Stitch Fix shows real-world use of image classification + recommendations",
            "Good discussion on human-AI collaboration",
            "Ask: Would you trust an algorithm or a human to pick your clothes?"
        ]
    },
]


# ============================================
# GENERATE NEW CASES
# ============================================

import os

os.makedirs("/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/student", exist_ok=True)
os.makedirs("/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/instructor", exist_ok=True)

for case in new_cases:
    # Student handout
    student_file = f"/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/student/{case['id']}.docx"
    create_student_handout(
        student_file,
        case['week'],
        case['topic'],
        case['title'],
        case['content_paragraphs'],
        case['key_facts'],
        case['questions'],
        case['primary_source']
    )

    # Instructor notes
    instructor_file = f"/sessions/bold-upbeat-keller/mnt/CLAUDE/case-handouts/instructor/{case['id']}_INSTRUCTOR.docx"
    create_instructor_notes(
        instructor_file,
        case['week'],
        case['topic'],
        case['title'],
        case['questions'],
        case['key_insight'],
        case.get('teaching_tips', [])
    )

print("\nNew cases generated successfully!")
print("- Week1B_HM: H&M Demand Forecasting")
print("- Week2B_Carvana: Carvana/CarMax Vehicle Pricing")
print("- Week5B_StitchFix: Stitch Fix AI Styling")
